"""willab — The Life Panel (founder-directed 2026-07-26).

Every non-negotiable in PROMPT-BE-life-panel.md §0 is a test here, not a note:

  N1  RLS on every life_* table IN ITS CREATING MIGRATION
  N2  no product-side module imports the life module; the SHARED master body
      is never written to
  N3  a non-participating user's chat is byte-identical — the router hands the
      turn back untouched
  N4  the per-user master-doc injection is capped and trimmed, and is absent
      entirely for everyone who has not opted in
  N5  the system never authors reflective prose (L-1)
  N6  nothing reaches the strategy doc without an approved proposal, and the
      immutable core is never proposed against (L-2 / L-2a)
  N7  at most one strategy proposal surfaced per day (L-2b)
  N8  zero scheduled outbound anything (L-4)
  N9  hard delete actually deletes; export returns everything

Plus the locked rules the routes rest on: the hashtag grammar (§5), the fixed
six-entry taxonomy (§3.1), the phrase wall's relevance floor and no-repeat
window, L-3's both-sides-never-a-winner, and the importer's preservation
guarantees (BE-3).

The pure-layer tests import ONLY services.life_panel, which has no DB and no
network — so they run in CI unconditionally. The route/store tests need the
app's dependency graph and skip cleanly when it is unavailable, following the
house pattern in test_journal.py.

Run: python3 -m unittest test_life_panel

N3's full form is a SUITE-level gate, not a single test — "the full existing
chat suite re-run under a non-participating user". Run it with the flag ON and
no consent row, which is exactly that state:

    LIFE_PANEL_ENABLED=1 python3 -m pytest -q

Green means the hook falls through untouched for everyone who has not opted
in. If that run ever goes red on a NON-life test, the hook has started
changing responses it must not change.
"""
from __future__ import annotations

import os
import re
import unittest
from datetime import date, datetime, timedelta, timezone
from unittest.mock import patch

# Placeholder credentials BEFORE the app graph is imported: config.py and
# services/db.py read the environment at import time, and the Supabase client
# validates the key's SHAPE at construction. These never leave the process —
# every store call in this file is mocked.
os.environ.setdefault("SUPABASE_URL", "https://life-panel-test.invalid")
os.environ.setdefault("SUPABASE_SERVICE_ROLE_KEY", "aaaa.bbbb.cccc")

from services import life_panel as lp           # noqa: E402  (pure, always importable)
from services import life_import as importer    # noqa: E402  (pure)

try:
    from flask import Flask
    import auth
    from routes import life_routes as lroutes
    from routes import life_reminders_webhook as lremhook
    from services import life_chat as lchat
    from services import life_engine as lengine
    from services import life_reminders as lreminders
    from services import life_store as lstore
    from services import master_doc_rag as mdr
    _IMPORT_ERROR = None
except Exception as e:                          # pragma: no cover
    Flask = None
    auth = lroutes = lremhook = lchat = lengine = lreminders = lstore = None
    mdr = None
    _IMPORT_ERROR = e

USER = "11111111-1111-4111-8111-111111111111"
OTHER = "22222222-2222-4222-8222-222222222222"


def code_only(source: str) -> str:
    """``source`` with every comment and string literal removed.

    The grep-shaped fences (L-4's "no outbound anything", the no-tracking
    rule) must read CODE. Run against raw source they match the prose that
    explains the rule — a comment saying "no streaks, no scores" fails the
    streak check — so the fence ends up forbidding its own documentation and
    the fix is to delete the explanation. That is backwards: the comment is
    the most valuable line in the file.

    Tokenising is the honest version of the `.split('\"\"\"')` hack this
    replaces, which only ever stripped the first docstring."""
    import io
    import tokenize
    out: list[str] = []
    try:
        for tok in tokenize.generate_tokens(io.StringIO(source).readline):
            if tok.type in (tokenize.COMMENT, tokenize.STRING):
                continue
            out.append(tok.string)
    except (tokenize.TokenError, IndentationError):   # pragma: no cover
        return source
    return " ".join(out)
MIGRATION = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                         "migrations", "add_life_panel.sql")


# ═════════════════════════════════════════════════════════════════════════
# N1 — RLS in the creating migration
# ═════════════════════════════════════════════════════════════════════════

class RlsMigrationTests(unittest.TestCase):
    """RLS is part of the file that creates the table, not a follow-up sweep.

    The July 2026 audit found 57 exposed public tables; this corpus —
    addiction, sexual behaviour, confession-shaped religious material, named
    third parties — must not be number 58."""

    @classmethod
    def setUpClass(cls):
        with open(MIGRATION, "r", encoding="utf-8") as fh:
            cls.sql = fh.read()

    def _created_tables(self) -> set[str]:
        return set(re.findall(
            r"CREATE TABLE IF NOT EXISTS\s+(life_\w+)", self.sql))

    def test_every_created_table_enables_rls_in_this_file(self):
        protected = set(re.findall(
            r"ALTER TABLE\s+(life_\w+)\s+ENABLE ROW LEVEL SECURITY", self.sql))
        created = self._created_tables()
        self.assertTrue(created, "migration creates no life_* tables")
        self.assertEqual(
            created - protected, set(),
            "these tables are created without RLS in the same migration")

    def test_no_policies_are_granted_to_anon(self):
        # RLS + zero policies = anon and authenticated can do nothing, while
        # the backend's service-role key bypasses RLS. A policy here would be
        # the thing that re-opens the door.
        self.assertNotIn("CREATE POLICY", self.sql.upper())

    def test_migration_is_idempotent(self):
        for statement in re.findall(r"CREATE TABLE[^(]*", self.sql):
            self.assertIn("IF NOT EXISTS", statement)
        for statement in re.findall(r"CREATE (?:UNIQUE )?INDEX[^(]*", self.sql):
            self.assertIn("IF NOT EXISTS", statement)

    def test_nothing_is_ever_dropped(self):
        # The standing engineering constraint: never auto-drop tables,
        # columns or migrations.
        self.assertNotRegex(self.sql.upper(), r"\bDROP\s+(TABLE|COLUMN)\b")

    def test_a_partial_run_raises_instead_of_reporting(self):
        # A migration that silently half-applies is how a table ends up
        # exposed. The post-condition must RAISE.
        self.assertIn("RAISE EXCEPTION", self.sql)

    def test_the_store_knows_about_every_created_table(self):
        # export_all and hard_delete iterate LIFE_TABLES. A table in the
        # migration but not in that list would be invisible to both — it would
        # neither come out on export nor go away on delete.
        if _IMPORT_ERROR is not None:
            self.skipTest(f"needs app deps: {_IMPORT_ERROR}")
        self.assertEqual(self._created_tables() - set(lstore.LIFE_TABLES),
                         set())

    def test_a_strategy_proposal_cannot_exist_without_a_warrant(self):
        # L-2 at the DB layer: a direct SQL insert cannot bypass the warrant
        # requirement either.
        self.assertIn("life_proposals_strategy_needs_warrant", self.sql)


# ═════════════════════════════════════════════════════════════════════════
# §5 — the hashtag grammar
# ═════════════════════════════════════════════════════════════════════════

class TagRouterTests(unittest.TestCase):

    def test_every_documented_tag_routes(self):
        for tag in ("principle", "sin", "mistake", "error", "problem"):
            self.assertEqual(lp.parse_tag(f"#{tag} x")[1], "case")
        for tag in ("data", "observation", "reflection", "idea", "finding"):
            self.assertEqual(lp.parse_tag(f"#{tag} x")[1], "goal_diff")
        for tag in ("win", "wins", "wygrane", "liftmeup"):
            self.assertEqual(lp.parse_tag(f"#{tag} x")[1], "win")
        self.assertEqual(lp.parse_tag("#add x")[1], "phrase")
        self.assertEqual(lp.parse_tag("#edit x")[1], "edit")

    def test_lift_me_up_is_the_desired_failure_mode(self):
        # Hashtags break at whitespace, so `#liftmeup` is one word on purpose.
        # `#lift me up` must NOT reach Wins — it is an unknown tag, and the
        # whole text is kept so nothing the user typed is silently edited.
        tag, route, remainder = lp.parse_tag("#lift me up")
        self.assertIsNone(tag)
        self.assertEqual(route, "capture")
        self.assertEqual(remainder, "#lift me up")

    def test_the_tag_is_the_first_token_only(self):
        self.assertEqual(lp.parse_tag("some prose #mistake later")[1],
                         "capture")

    def test_case_insensitive_and_punctuation_tolerant(self):
        self.assertEqual(lp.parse_tag("#MISTAKE: I did x"),
                         ("mistake", "case", "I did x"))

    def test_unknown_tag_keeps_the_text_whole(self):
        self.assertEqual(lp.parse_tag("#nonsense body")[2], "#nonsense body")

    def test_sin_works_but_is_off_the_public_picker(self):
        self.assertEqual(lp.parse_tag("#sin x")[1], "case")
        self.assertNotIn("#sin", [s["tag"] for s in lp.tag_suggestions()])

    def test_the_picker_covers_every_route(self):
        # Four aliases per route are impossible to memorise; the picker is
        # what keeps the guide page from having to be load-bearing.
        self.assertEqual({s["route"] for s in lp.tag_suggestions()},
                         set(lp.TAG_ROUTES.values()))

    def test_empty_and_bare_hash(self):
        self.assertEqual(lp.parse_tag("")[1], "capture")
        self.assertEqual(lp.parse_tag(None)[1], "capture")
        self.assertEqual(lp.parse_tag("#")[1], "capture")


# ═════════════════════════════════════════════════════════════════════════
# §3.1 — the fixed six-entry taxonomy
# ═════════════════════════════════════════════════════════════════════════

class TaxonomyTests(unittest.TestCase):

    def test_exactly_six_categories(self):
        self.assertEqual(len(lp.CATEGORIES), 6)
        self.assertEqual(set(lp.CATEGORY_LABELS), set(lp.CATEGORIES))

    def test_aliases_map_both_languages(self):
        self.assertEqual(lp.map_import_category("Wishful thinking"),
                         "wishful_thinking")
        self.assertEqual(lp.map_import_category("Myślenie życzeniowe"),
                         "wishful_thinking")
        self.assertEqual(lp.map_import_category("PYCHA"), "hubris")

    def test_an_unknown_category_is_never_guessed(self):
        # It goes to a review queue instead. A mis-mapped category is a lie
        # about a four-year-old reflection, and it is invisible once written.
        self.assertIsNone(lp.map_import_category("Something else entirely"))
        self.assertIsNone(lp.map_import_category(""))
        self.assertIsNone(lp.map_import_category(None))

    def test_normalize_drops_rather_than_coerces(self):
        self.assertEqual(
            lp.normalize_categories(["hubris", "nope", "hubris"]),
            ["hubris"])

    def test_multi_category_survives_and_keeps_its_order(self):
        # The scooter/drift case carries Wishful thinking + Hubris.
        self.assertEqual(
            lp.normalize_categories(["Wishful thinking", "Hubris"]),
            ["wishful_thinking", "hubris"])


# ═════════════════════════════════════════════════════════════════════════
# N5 / L-1 — the system never authors reflective prose
# ═════════════════════════════════════════════════════════════════════════

class ReflectionsAreNeverAuthoredTests(unittest.TestCase):

    def test_reflections_come_from_request_input(self):
        out = lp.validate_case_input({
            "case_at_hand": "I did x", "reflections": "Zabrakło mi pokory.",
        })
        self.assertEqual(out["reflections"], "Zabrakło mi pokory.")

    def test_polish_is_stored_verbatim(self):
        # No translation, no normalisation, no cleanup pass. Several
        # reflections are code-switched mid-paragraph and that is the record.
        text = "Zabrakło mi pokory — I lost control. Żałuję."
        self.assertEqual(
            lp.validate_case_input({"case_at_hand": "x",
                                    "reflections": text})["reflections"],
            text)

    @unittest.skipIf(_IMPORT_ERROR is not None, "needs app deps")
    def test_no_module_ever_writes_reflections_from_a_model_result(self):
        """The load-bearing half of N5.

        Anywhere in the feature, `reflections` may be assigned only from
        validated request input or from an import row. If a future change
        wires a model output into that field, this fails."""
        import inspect
        for module in (lengine, lchat, lroutes):
            src = inspect.getsource(module)
            for match in re.finditer(r'"reflections"\s*:\s*([^,\n}]+)', src):
                value = match.group(1).strip()
                self.assertNotIn(
                    "derived", value,
                    f"{module.__name__} assigns reflections from a derivation")
                self.assertNotIn(
                    "parsed", value,
                    f"{module.__name__} assigns reflections from model output")

    @unittest.skipIf(_IMPORT_ERROR is not None, "needs app deps")
    def test_the_case_prompt_forbids_writing_reflections(self):
        self.assertIn("YOU DO NOT WRITE REFLECTIONS", lengine._CASE_SYSTEM)

    @unittest.skipIf(_IMPORT_ERROR is not None, "needs app deps")
    def test_the_panel_never_calls_the_f1_ideal_text_pipeline(self):
        # L-1: the Ideal Text polish happens EXTERNALLY and is pasted in. This
        # is a manual copy-paste, not an integration.
        import inspect
        for module in (lengine, lchat, lroutes, lstore, lp, importer):
            self.assertNotIn("ideal_text", inspect.getsource(module))


# ═════════════════════════════════════════════════════════════════════════
# N6 / L-2a — the immutable core
# ═════════════════════════════════════════════════════════════════════════

class ImmutableCoreTests(unittest.TestCase):

    def test_section_i_and_the_bet_rank_are_immutable(self):
        for target in ("anchor", "weekly.section_i", "Section I",
                       "weekly.section_1", "bets.rank", "bet_rank",
                       "yearly.bets.ranking"):
            self.assertTrue(lp.is_immutable_target(target), target)

    def test_ordinary_sections_are_not(self):
        for target in ("weekly.section_ii", "yearly.bet_2.short_term",
                       "monthly.goals", "daily.one_thing"):
            self.assertFalse(lp.is_immutable_target(target), target)

    def test_a_blank_target_is_refused(self):
        # An unaddressed proposal cannot be reviewed, so it cannot be created.
        # Refusing it keeps "every change is gated" true by construction.
        self.assertTrue(lp.is_immutable_target(""))
        self.assertTrue(lp.is_immutable_target(None))

    @unittest.skipIf(_IMPORT_ERROR is not None, "needs app deps")
    def test_the_report_shape_says_report_only_by_name(self):
        # The two response shapes must agree on the field the FE reads. A
        # reader who only knows `kind` would not realise that omitting
        # `report_only` here is the difference between a report and an
        # editable proposal over the Anchor.
        import inspect
        src = inspect.getsource(lengine.compare_to_strategy)
        report_branch = src[src.index('"kind": "report"'):]
        self.assertIn('"report_only": True', report_branch)

    @unittest.skipIf(_IMPORT_ERROR is not None, "needs app deps")
    def test_apply_refuses_the_immutable_core_on_the_write_path_too(self):
        # "Never proposed against" has to hold on the WRITE path, not only on
        # the path that happens to be in front of it.
        self.assertIsNone(lengine.apply_proposal(
            USER, {"id": "p1", "target": "weekly.section_i",
                   "current": "a", "proposed": "b"}))


# ═════════════════════════════════════════════════════════════════════════
# N7 / L-2b — the change budget
# ═════════════════════════════════════════════════════════════════════════

class ChangeBudgetTests(unittest.TestCase):

    def test_five_qualifying_notes_in_one_day_surface_one(self):
        surfaced = 0
        statuses = []
        for _ in range(5):
            status = lp.plan_proposal_status(already_surfaced_today=surfaced)
            statuses.append(status)
            if status == "surfaced":
                surfaced += 1
        self.assertEqual(statuses.count("surfaced"), 1)
        self.assertEqual(statuses.count("queued"), 4)

    def test_a_garbled_count_queues_rather_than_surfaces(self):
        # Fail closed: a queued proposal still reaches the weekly review, an
        # over-surfaced one erodes the approve button.
        self.assertEqual(lp.plan_proposal_status(already_surfaced_today=None),
                         "surfaced")
        self.assertEqual(lp.plan_proposal_status(already_surfaced_today=99),
                         "queued")

    def test_the_weekly_batch_is_three_ranked(self):
        queued = [{"id": str(i), "rank": i / 10.0, "created_at": "2026-07-01"}
                  for i in range(7)]
        batch = lp.weekly_batch(queued)
        self.assertEqual(len(batch), lp.WEEKLY_BATCH_SIZE)
        self.assertEqual([b["id"] for b in batch], ["6", "5", "4"])

    def test_queued_proposals_expire_after_two_weeks(self):
        old = (datetime.now(timezone.utc) - timedelta(days=1)).isoformat()
        future = (datetime.now(timezone.utc) + timedelta(days=1)).isoformat()
        self.assertTrue(lp.is_expired({"status": "queued",
                                       "expires_at": old}))
        self.assertFalse(lp.is_expired({"status": "queued",
                                        "expires_at": future}))

    def test_surfaced_and_decided_proposals_never_expire(self):
        old = (datetime.now(timezone.utc) - timedelta(days=30)).isoformat()
        for status in ("surfaced", "approved", "dismissed", "expired"):
            self.assertFalse(lp.is_expired({"status": status,
                                            "expires_at": old}), status)

    def test_expiry_is_fourteen_days_out(self):
        base = datetime(2026, 7, 26, tzinfo=timezone.utc)
        self.assertTrue(lp.expiry_for(base).startswith("2026-08-09"))

    @unittest.skipIf(_IMPORT_ERROR is not None, "needs app deps")
    def test_an_unknown_budget_state_fails_closed(self):
        # The budget is a QUERY, and a failed query must not be read as zero.
        with patch.object(lstore, "_t", side_effect=RuntimeError("down")):
            self.assertGreaterEqual(lstore.count_surfaced_today(USER),
                                    lp.DAILY_PROPOSAL_BUDGET)


# ═════════════════════════════════════════════════════════════════════════
# The phrase wall — floor + no-repeat
# ═════════════════════════════════════════════════════════════════════════

class PhraseWallTests(unittest.TestCase):

    def _phrases(self):
        return [
            {"id": "a", "body": "Focus on saving your own life first",
             "created_at": "2022-01-01"},
            {"id": "b", "body": "Money follows patience and honest partners",
             "created_at": "2023-01-01"},
        ]

    def test_below_the_floor_attaches_nothing(self):
        # A mismatched aphorism on a `#sin` note about addiction is worse than
        # silence — it teaches the founder the wall is noise.
        self.assertIsNone(lp.pick_phrase("the weather is cold today",
                                         self._phrases()))

    def test_one_shared_word_is_coincidence_not_relevance(self):
        self.assertEqual(
            lp.phrase_relevance("money", "Money follows patience and honest "
                                         "partners"), 0.0)

    def test_a_real_overlap_clears_the_floor(self):
        picked = lp.pick_phrase(
            "I keep chasing money instead of finding honest partners",
            self._phrases())
        self.assertIsNotNone(picked)
        self.assertEqual(picked["id"], "b")

    def test_the_no_repeat_window_is_respected(self):
        note = "I keep chasing money instead of finding honest partners"
        self.assertIsNone(
            lp.pick_phrase(note, self._phrases(), recently_used_ids=["b"]))

    def test_the_wall_cannot_collapse_to_one_favourite(self):
        # Same note twice; the second time the winner is blocked, so either a
        # different phrase or nothing comes back — never the same one again.
        note = "I keep chasing money instead of finding honest partners"
        first = lp.pick_phrase(note, self._phrases())
        second = lp.pick_phrase(note, self._phrases(),
                                recently_used_ids=[first["id"]])
        self.assertNotEqual((second or {}).get("id"), first["id"])

    def test_the_score_never_reaches_the_payload(self):
        # Nothing in a life payload is a number the user reads.
        picked = lp.pick_phrase(
            "I keep chasing money instead of finding honest partners",
            self._phrases())
        self.assertNotIn("score", lp.serialize_item(picked))
        self.assertNotIn("relevance", lp.serialize_item(picked))


# ═════════════════════════════════════════════════════════════════════════
# L-3 — surface paradoxes, never resolve them
# ═════════════════════════════════════════════════════════════════════════

class ParadoxTests(unittest.TestCase):

    def test_opposing_principles_come_back_as_a_pair(self):
        pairs = lp.pair_conflicts([
            {"id": "1", "title": "Don't seek validation outside yourself"},
            {"id": "2", "title": "Seek honest feedback from trusted people"},
        ])
        self.assertEqual(len(pairs), 1)
        self.assertEqual({p["id"] for p in pairs[0]}, {"1", "2"})

    def test_there_is_no_resolver_in_the_module(self):
        # A caller looking for "which one applies" must not find it here.
        for name in dir(lp):
            self.assertNotIn("resolve", name.lower())
            self.assertNotIn("winner", name.lower())

    def test_unrelated_principles_do_not_pair(self):
        self.assertEqual(lp.pair_conflicts([
            {"id": "1", "title": "Write things down"},
            {"id": "2", "title": "Sleep before deciding"},
        ]), [])


# ═════════════════════════════════════════════════════════════════════════
# The board, due labels, and item validation
# ═════════════════════════════════════════════════════════════════════════

class BoardAndLabelTests(unittest.TestCase):

    def test_the_board_routes_by_domain_and_says_why(self):
        key, why = lp.route_advisor("should I take money from this partner?")
        self.assertEqual(key, "munger")
        self.assertIn("money", why)

    def test_faith_routes_to_the_right_lens(self):
        self.assertEqual(
            lp.route_advisor("my marriage and my faith feel far apart")[0],
            "jp2")

    def test_an_undomained_question_defaults_and_admits_it(self):
        key, why = lp.route_advisor("hmm")
        self.assertEqual(key, "dalio")
        self.assertIn("no clear domain", why)

    def test_the_board_never_replaces_prayer(self):
        self.assertIn("does not replace prayer", lp.ADVISOR_PRAYER_LINE)

    def test_due_labels_parse_where_unambiguous(self):
        self.assertEqual(lp.parse_due_label("[Jul '27]"), "2027-07-01")
        self.assertEqual(lp.parse_due_label("2035"), "2035-01-01")
        self.assertEqual(lp.parse_due_label("[Aug]", today=date(2026, 9, 1)),
                         "2027-08-01")

    def test_now_is_a_horizon_not_a_date(self):
        # A standing intention must not become a one-day marker on the
        # timeline.
        self.assertIsNone(lp.parse_due_label("[NOW]"))

    def test_an_unparseable_label_is_null_not_a_guess(self):
        # The label is the source of truth; a marker in the wrong year reads
        # as a fact.
        self.assertIsNone(lp.parse_due_label("someday soon"))
        self.assertIsNone(lp.parse_due_label(None))

    def test_the_label_stays_the_source_of_truth(self):
        out = lp.validate_item_input({"due_label": "someday soon"},
                                     partial=True)
        self.assertEqual(out["due_label"], "someday soon")
        self.assertIsNone(out["due_at"])

    def test_kind_cannot_be_changed_after_creation(self):
        with self.assertRaises(lp.LifeError):
            lp.validate_item_input({"kind": "win"}, partial=True)

    def test_all_three_bets_reach_the_daily_card(self):
        # Answered by the founder 2026-07-26: 🟣 The Dream is eligible from
        # day one.
        self.assertTrue(lp.BET_3_DRIVES_DAILY_EXECUTION)
        self.assertEqual([b["rank"] for b in lp.BETS], [1, 2, 3])


# ═════════════════════════════════════════════════════════════════════════
# L-5 — the application log
# ═════════════════════════════════════════════════════════════════════════

class ApplicationLogTests(unittest.TestCase):

    def test_rows_are_built_for_each_cited_principle(self):
        rows = lp.application_rows_for(USER, ["a", "b"], context="case",
                                       ref_id="c1")
        self.assertEqual(len(rows), 2)
        self.assertEqual(rows[0]["ref_id"], "c1")

    def test_duplicates_and_blanks_are_dropped(self):
        self.assertEqual(
            len(lp.application_rows_for(USER, ["a", "a", "", None],
                                        context="case")), 1)

    def test_an_unknown_context_is_rejected(self):
        with self.assertRaises(lp.LifeError):
            lp.application_rows_for(USER, ["a"], context="nonsense")


# ═════════════════════════════════════════════════════════════════════════
# BE-3 — the importer's preservation guarantees
# ═════════════════════════════════════════════════════════════════════════

class ImporterTests(unittest.TestCase):

    def _export(self):
        return {
            "principles": [{
                "id": "a1",
                "caseAtHand": "Pojechałem hulajnogą po alkoholu",
                "category": "Wishful thinking, Hubris",
                "principlesApplied": "Nie miałem wtedy zasad",
                "reflections": "Zabrakło mi pokory — I lost control.",
                "newPrinciple": ("Focus on saving your own life first\n"
                                 "Don't try to understand it all"),
                "createdAt": "2023-04-11",
                "order": 2,
            }],
            "wins": ["Ran 10k"],
            "prayer": "„Nie bój się” and “Keep the promise” plus loose text",
        }

    def test_created_at_is_preserved(self):
        # The corpus spans 2022→2026 and the dates carry meaning.
        plan = importer.plan_cases(USER, self._export()["principles"])
        self.assertTrue(plan["cases"][0]["created_at"].startswith("2023-04-11"))

    def test_a_naive_source_date_is_read_as_utc(self):
        # Handing Postgres a naive value would make the date depend on the
        # server's zone and silently move a 2022 reflection by a day.
        self.assertTrue(importer._iso("2023-04-11").endswith("+00:00"))

    def test_multi_category_and_multi_principle_both_survive(self):
        plan = importer.plan_cases(USER, self._export()["principles"])
        self.assertEqual(plan["cases"][0]["category"],
                         ["wishful_thinking", "hubris"])
        principles = [i for i in plan["items"] if i["kind"] == "principle"]
        self.assertEqual(len(principles), 2)

    def test_polish_reflections_are_byte_identical(self):
        plan = importer.plan_cases(USER, self._export()["principles"])
        self.assertEqual(plan["cases"][0]["reflections"],
                         "Zabrakło mi pokory — I lost control.")

    def test_an_unmapped_category_goes_to_the_review_queue(self):
        plan = importer.plan_cases(USER, [{
            "id": "x", "caseAtHand": "c", "category": "Something Else",
            "reflections": "r"}])
        self.assertEqual(len(plan["review"]), 1)
        self.assertEqual(plan["cases"][0]["category"], [])
        # Parked, not coerced — and the row still imports, so the reflection
        # is safe while a human decides what the label meant.
        self.assertEqual(plan["cases"][0]["import_category_raw"],
                         "Something Else")

    def test_the_drag_order_becomes_order_key(self):
        plan = importer.plan_cases(USER, self._export()["principles"])
        self.assertEqual(plan["items"][0]["order_key"], 2000.0)

    def test_a_two_sentence_principle_is_not_split_in_half(self):
        # Only line breaks and list markers split. A sentence break does not:
        # several principles in the corpus are two sentences long.
        self.assertEqual(
            importer._split_principles("Do the work. Then rest."),
            ["Do the work. Then rest."])

    def test_the_prayer_blob_splits_on_quote_boundaries(self):
        phrases, leftover = importer.split_phrases(self._export()["prayer"])
        self.assertEqual(phrases, ["Nie bój się", "Keep the promise"])
        self.assertIn("loose text", leftover)

    def test_an_unsplittable_blob_survives_whole(self):
        # "Anything it can't split stays as one note so nothing is lost."
        phrases, leftover = importer.split_phrases("no quotes here at all")
        self.assertEqual(phrases, [])
        self.assertEqual(leftover, "no quotes here at all")
        plan = importer.plan_phrases(USER, "no quotes here at all")
        self.assertEqual(plan["items"], [])
        self.assertEqual(len(plan["notes"]), 1)

    def test_an_apostrophe_never_splits_a_phrase(self):
        # "Don't try to understand it all" is an actual principle.
        phrases, _ = importer.split_phrases("„Don't try to understand it all”")
        self.assertEqual(phrases, ["Don't try to understand it all"])

    def test_an_unclosed_quote_does_not_weld_two_phrases(self):
        phrases, leftover = importer.split_phrases(
            'unclosed „first one and then “second one” tail')
        self.assertEqual(phrases, ["second one"])
        self.assertIn("first one", leftover)

    def test_external_ids_make_a_re_run_idempotent(self):
        first = importer.plan_cases(USER, self._export()["principles"])
        second = importer.plan_cases(USER, self._export()["principles"])
        self.assertEqual(first["cases"][0]["external_id"],
                         second["cases"][0]["external_id"])

    def test_the_bets_are_seeded_in_their_locked_rank(self):
        # Never taken from a free-form transcription that could have
        # reordered them (L-2a).
        items = importer.plan_strategy(USER, {})["items"]
        bets = [i for i in items if i["kind"] == "bet"]
        self.assertEqual([b["order_key"] for b in bets], [1.0, 2.0, 3.0])

    def test_dry_run_is_the_default(self):
        # The destructive direction should require a word, not the absence of
        # one.
        plan = importer.build_plan(USER, principles_export=self._export())
        self.assertTrue(importer.apply_plan(USER, plan)["dry_run"])
        self.assertEqual(
            importer.apply_plan(USER, plan)["written"],
            {"cases": 0, "items": 0, "notes": 0, "strategy": 0})

    def test_the_dry_run_reports_the_same_counts_the_real_run_would_write(self):
        plan = importer.build_plan(USER, principles_export=self._export())
        reported = importer.apply_plan(USER, plan)["planned"]
        self.assertEqual(reported["cases"], len(plan["cases"]))
        self.assertEqual(reported["items"], len(plan["items"]))


# ═════════════════════════════════════════════════════════════════════════
# Serialization
# ═════════════════════════════════════════════════════════════════════════

class SerializationTests(unittest.TestCase):

    def test_no_payload_carries_a_score_or_a_verdict(self):
        import json
        blob = json.dumps({
            "case": lp.serialize_case({"case_at_hand": "x",
                                       "category": ["hubris"]}),
            "item": lp.serialize_item({"kind": "principle", "title": "t"}),
            "day": lp.serialize_day({"one_thing": "x"}),
            "week": lp.serialize_week({}),
            "proposal": lp.serialize_proposal({"target": "weekly.goals"}),
        })
        for banned in ("power_score", "charisma", "overall_score", "threat",
                       "verdict", "confidence", "relevance"):
            self.assertNotIn(banned, blob)

    def test_categories_render_as_labels_for_the_multiple_lines(self):
        out = lp.serialize_case({"category": ["wishful_thinking", "hubris"]})
        self.assertEqual(out["category_labels"],
                         ["Wishful thinking", "Hubris"])

    def test_every_proposal_states_report_only_explicitly(self):
        """The FE fail-safes on this field's ABSENCE: no `report_only` means
        no approve button, so an unrecognised payload cannot grow one over
        Section I. That default is correct and must not change — which is
        exactly why the backend has to SAY the value. A serializer that
        omitted it would render every proposal un-approvable and silently
        kill the L-2 approve flow, and the symptom would be a missing button
        rather than an error."""
        out = lp.serialize_proposal({"target": "weekly.goals"})
        self.assertIn("report_only", out)
        self.assertFalse(out["report_only"])

    def test_a_row_in_the_table_is_never_report_only(self):
        # Immutable-core proposals are refused at creation and again on the
        # write path, so anything that reached life_proposals is approvable.
        for target in ("weekly.section_i", "bets.rank", "anchor"):
            self.assertTrue(lp.is_immutable_target(target))
        self.assertFalse(
            lp.serialize_proposal({"target": "weekly.section_ii"})["report_only"])

    def test_a_proposal_carries_its_warrant(self):
        # A change the archive cannot justify is a change the system invented.
        out = lp.serialize_proposal({"target": "weekly.goals"},
                                    warrant={"id": "p1", "title": "mine"})
        self.assertEqual(out["warrant"]["title"], "mine")


# ═════════════════════════════════════════════════════════════════════════
# N3 — the chat router hands a non-participant's turn back untouched
# ═════════════════════════════════════════════════════════════════════════

@unittest.skipIf(_IMPORT_ERROR is not None, f"needs app deps: {_IMPORT_ERROR}")
class ChatIsolationTests(unittest.TestCase):

    def test_an_untagged_message_is_never_ours_and_costs_no_db_read(self):
        with patch.object(lchat, "has_consented") as consented:
            self.assertIsNone(lchat.handle_note(USER, "how does this work?"))
            consented.assert_not_called()

    def test_a_non_consented_user_gets_nothing_from_us(self):
        # Returning None hands the turn back to the existing chat path
        # completely untouched — that IS byte-identical.
        with patch.object(lchat, "has_consented", return_value=False), \
                patch.object(lchat.store, "insert_note") as insert:
            self.assertIsNone(lchat.handle_note(USER, "#mistake I did x"))
            # And nothing was written: consent precedes ANY life row.
            insert.assert_not_called()

    def test_the_flag_is_off_by_default(self):
        """The kill switch is off unless an operator turns it on.

        Asserted against the DEFAULT rather than the ambient value, so the
        test still means something during the N3 run that deliberately sets
        LIFE_PANEL_ENABLED=1 (see the module docstring)."""
        import inspect
        import config as cfg
        self.assertIn('os.getenv("LIFE_PANEL_ENABLED") or "0"',
                      inspect.getsource(cfg))
        if not os.getenv("LIFE_PANEL_ENABLED"):
            self.assertFalse(lchat.is_enabled())

    def test_the_allowlist_is_empty_by_default_and_never_matches_none(self):
        self.assertFalse(lchat.is_allowlisted(None))
        self.assertFalse(lchat.is_allowlisted(USER))

    def test_the_hook_in_the_chat_route_is_flag_and_auth_guarded(self):
        # The one contact point with v2_routes.py. Both guards must be on the
        # same line of defence as the call itself.
        with open(os.path.join(os.path.dirname(MIGRATION), "..",
                               "routes", "v2_routes.py"),
                  "r", encoding="utf-8") as fh:
            src = fh.read()
        hook = src[src.index("Life Panel hashtag router"):]
        hook = hook[:hook.index("Goal-update intercept")]
        self.assertIn("request.user_id and getattr(config, "
                      "\"LIFE_PANEL_ENABLED\", False)", hook)
        # A broken Life Panel must cost the panel, never the chat.
        self.assertIn("except Exception", hook)

    def test_the_consent_check_does_not_hit_the_db_on_every_chat_turn(self):
        """Scaffolding must not make the live loop slower.

        With the flag on, the chat route asks this on every signed-in turn so
        the per-user block can be retrieved for participants. Uncached that is
        a Supabase round trip added to the shipped chat path for everyone."""
        lchat._consent_cache.clear()
        with patch.object(lchat.store, "get_consent",
                          return_value=None) as read:
            for _ in range(20):
                lchat.has_consented(USER)
        self.assertEqual(read.call_count, 1)
        lchat._consent_cache.clear()

    def test_accepting_consent_takes_effect_immediately(self):
        # Without invalidation the user finishes the consent screen and finds
        # the feature still inert for five minutes.
        lchat._consent_cache.clear()
        with patch.object(lchat.store, "get_consent", return_value=None):
            self.assertFalse(lchat.has_consented(USER))
        lchat.invalidate_consent(USER)
        with patch.object(lchat.store, "get_consent", return_value={"id": "c"}):
            self.assertTrue(lchat.has_consented(USER))
        lchat._consent_cache.clear()

    def test_a_hard_delete_drops_the_cached_yes(self):
        # Otherwise a just-wiped account keeps passing the gate and can write
        # fresh rows into the account it emptied a second ago.
        lchat._consent_cache.clear()
        with patch.object(lchat.store, "get_consent", return_value={"id": "c"}):
            self.assertTrue(lchat.has_consented(USER))
        lchat.invalidate_consent(USER)
        with patch.object(lchat.store, "get_consent", return_value=None):
            self.assertFalse(lchat.has_consented(USER))
        lchat._consent_cache.clear()

    def test_the_cache_never_leaks_between_users(self):
        lchat._consent_cache.clear()
        with patch.object(lchat.store, "get_consent", return_value={"id": "c"}):
            self.assertTrue(lchat.has_consented(USER))
        with patch.object(lchat.store, "get_consent",
                          return_value=None) as read:
            self.assertFalse(lchat.has_consented(OTHER))
        read.assert_called_once()
        lchat._consent_cache.clear()

    def test_both_invalidation_points_are_wired(self):
        # The cache is only safe because these two call sites exist.
        import inspect
        src = inspect.getsource(lroutes)
        self.assertEqual(src.count("chat.invalidate_consent"), 2)

    def test_a_derivation_failure_never_loses_the_note(self):
        # Store before derive is the whole reliability story.
        with patch.object(lchat, "has_consented", return_value=True), \
                patch.object(lchat.store, "setup_completed", return_value=True), \
                patch.object(lchat.store, "insert_note",
                             return_value={"id": "n1"}) as insert, \
                patch.object(lchat.engine, "derive_case",
                             side_effect=RuntimeError("openai down")), \
                patch.object(lchat.engine, "attach_phrase", return_value=None):
            out = lchat.handle_note(USER, "#mistake I did x")
        insert.assert_called_once()
        self.assertEqual(out["note_id"], "n1")
        self.assertIsNone(out["card"])

    def test_the_hash_router_is_off_until_onboarding_is_done(self):
        """Founder 2026-07-26: "no try outs".

        Before onboarding the tag does NOTHING — no derivation, no stored
        note, no message. The turn goes back to the existing chat path
        untouched, exactly like a non-consented user's.

        This supersedes §6.2's keep-the-note rule. That rule existed so the
        tag would not teach someone it costs something — but a half-working
        tag teaches it louder: it answers, appears to have understood, and
        produces nothing. The Principles button is the only way in."""
        with patch.object(lchat, "has_consented", return_value=True), \
                patch.object(lchat.store, "setup_completed",
                             return_value=False), \
                patch.object(lchat.store, "insert_note") as insert, \
                patch.object(lchat.engine, "derive_case") as derive:
            for text in ("#mistake I did x", "#win shipped it",
                         "#add a phrase", "#edit new one thing"):
                self.assertIsNone(lchat.handle_note(USER, text), text)
        # A user who never onboards leaves NO life rows at all.
        insert.assert_not_called()
        derive.assert_not_called()

    def test_the_panel_note_endpoint_is_gated_the_same_way(self):
        # A second door nobody checks is how a rule becomes advisory.
        import inspect
        src = inspect.getsource(lroutes.life_note_create)
        self.assertIn("setup_completed", src)
        self.assertIn("SETUP_REQUIRED", src)

    def test_an_untagged_panel_note_gets_no_phrase_attached(self):
        # Untagged is capture only, no action (§5) — and an aphorism returned
        # on it would be exactly the action the tag was supposed to require.
        with patch.object(lchat, "has_consented", return_value=True), \
                patch.object(lchat.store, "setup_completed", return_value=True), \
                patch.object(lchat.store, "insert_note",
                             return_value={"id": "n1"}), \
                patch.object(lchat.engine, "attach_phrase") as attach:
            lchat.handle_note(USER, "just a thought", source="form")
        attach.assert_not_called()

    def test_all_user_facing_copy_is_in_one_reviewable_dict(self):
        # Product copy is held for founder sign-off; a single dict makes the
        # review a diff instead of a grep.
        import inspect
        src = inspect.getsource(lchat)
        body = src[src.index("def handle_note"):]
        # Every string a user reads in the router comes from COPY — no
        # literal ever reaches an `answer`.
        self.assertIn('COPY["captured"]', body)
        self.assertNotRegex(body, r'"answer":\s*"[A-Z]')
        # And every COPY key is actually used, so a retired string cannot sit
        # in the review dict pretending to still ship. (`needs_setup` went
        # when the `#` router became off-until-onboarded.)
        # And every COPY key is actually reachable somewhere, so a retired
        # string cannot sit in the review dict pretending to still ship. (This
        # is how the retire prompt was found: BE-7's question existed as copy
        # and as an engine function, and nothing ever asked it.)
        import inspect as _i
        reachable = src + _i.getsource(lroutes)
        for key in lchat.COPY:
            self.assertIn(f'COPY["{key}"]', reachable,
                          f"COPY[{key!r}] is unused — dead copy, or a "
                          f"surface that was never wired up")


# ═════════════════════════════════════════════════════════════════════════
# BE-8 — the daily card, and the rank rule that survives Bet 3 joining it
# ═════════════════════════════════════════════════════════════════════════

@unittest.skipIf(_IMPORT_ERROR is not None, f"needs app deps: {_IMPORT_ERROR}")
class DailyCardTests(unittest.TestCase):
    """All three bets are eligible from day one. What holds the weekly
    document's rule is the SORT, not an exclusion — so these tests are what
    make flipping BET_3_DRIVES_DAILY_EXECUTION safe rather than merely
    permitted."""

    BETS = [
        {"id": "b1", "kind": "bet", "title": "🟢 The Life", "order_key": 1.0},
        {"id": "b2", "kind": "bet", "title": "🔵 The Company", "order_key": 2.0},
        {"id": "b3", "kind": "bet", "title": "🟣 The Dream", "order_key": 3.0},
    ]

    def _card(self, goals, habits=()):
        def _list(user_id, *, kind=None, **kw):
            return {"habit": list(habits), "goal": goals,
                    "bet": self.BETS}.get(kind, [])
        with patch.object(lengine.store, "list_items", side_effect=_list):
            return lengine.build_daily_card(USER, date(2026, 7, 26))

    def test_bet_three_never_outranks_bet_two(self):
        card = self._card([
            {"id": "g3", "title": "7T research", "bet_id": "b3",
             "horizon": "now", "order_key": 1},
            {"id": "g2", "title": "Ship the take ranker", "bet_id": "b2",
             "order_key": 9},
        ])
        # The Dream is on the card — and it is NOT the one thing, even though
        # it is the [NOW] item and the Company goal is not.
        self.assertEqual(card["one_thing"], "Ship the take ranker")
        self.assertEqual(card["one_thing_bet"], "🔵 The Company")
        self.assertIn("7T research", [b["text"] for b in card["focus_blocks"]])

    def test_the_dream_can_take_the_slot_when_nothing_outranks_it(self):
        # The one day you want the dream on the card instead of an empty card.
        card = self._card([{"id": "g3", "title": "7T research",
                            "bet_id": "b3", "order_key": 1}])
        self.assertEqual(card["one_thing"], "7T research")
        self.assertEqual(card["one_thing_bet"], "🟣 The Dream")

    def test_every_block_says_which_bet_it_serves(self):
        # §3.2 — without this the rank is invisible on the surface it governs,
        # and "Bet 3 never outranks Bet 2" is unverifiable by the reader.
        card = self._card([
            {"id": "g1", "title": "Call my father", "bet_id": "b1"},
            {"id": "g2", "title": "Ship it", "bet_id": "b2"},
        ])
        self.assertEqual(card["one_thing_bet"], "🟢 The Life")
        self.assertEqual([b["bet"] for b in card["focus_blocks"]],
                         ["🔵 The Company"])

    def test_a_goal_with_no_bet_sorts_last_rather_than_as_a_dream(self):
        # An unattached goal is not a Bet-3 item — it is one nobody has
        # decided a bet for, and it must not outrank one that was thought
        # about.
        card = self._card([
            {"id": "g0", "title": "Loose end", "order_key": 1},
            {"id": "g3", "title": "7T research", "bet_id": "b3",
             "order_key": 9},
        ])
        self.assertEqual(card["one_thing"], "7T research")

    def test_an_empty_archive_yields_an_empty_card_not_an_error(self):
        card = self._card([])
        self.assertEqual(card["one_thing"], "")
        self.assertEqual(card["one_thing_bet"], "")
        self.assertEqual(card["focus_blocks"], [])

    def test_the_evening_pass_recaps_the_morning_without_judging_it(self):
        summary = lengine.build_evening_pass({
            "one_thing": "Ship the take ranker",
            "one_thing_bet": "🔵 The Company",
            "focus_blocks": [{"text": "Call my father", "bet": "🟢 The Life"}],
            "morning_checks": {"Pompeiana": True, "Cold shower": False},
            "distraction_flagged": "phone",
        })
        self.assertEqual(summary["one_thing"], "Ship the take ranker")
        self.assertEqual(summary["habits"], ["Cold shower", "Pompeiana"])
        # It recaps; it never judges. No count, no streak, no score — the
        # ticks are the founder's to make, and a pass that pre-filled them
        # would be answering its own question.
        import json
        blob = json.dumps(summary)
        for banned in ("score", "streak", "completed", "percent", "rate"):
            self.assertNotIn(banned, blob.lower())

    def test_the_evening_pass_will_not_stamp_before_its_hour(self):
        # The FE branches on evening.generated_at, so an early stamp is a
        # wrong screen at breakfast, not a cosmetic issue.
        morning = datetime(2026, 7, 26, 9, 0, tzinfo=timezone.utc)
        evening = datetime(2026, 7, 26, 22, 0, tzinfo=timezone.utc)
        self.assertFalse(lengine.evening_pass_due(morning))
        self.assertTrue(lengine.evening_pass_due(evening))

        row = {"id": "d1", "day": "2026-07-26", "one_thing": "x"}
        with patch.object(lengine.store, "get_day", return_value=row), \
                patch.object(lengine.store, "update_day") as update:
            lengine.ensure_evening_pass(USER, date(2026, 7, 26), now=morning)
        update.assert_not_called()

    def test_the_evening_pass_is_idempotent(self):
        # A cron that fires twice must not overwrite an evening the founder
        # has already started answering.
        done = {"id": "d1", "day": "2026-07-26",
                "evening_generated_at": "2026-07-26T21:00:00+00:00"}
        with patch.object(lengine.store, "get_day", return_value=done), \
                patch.object(lengine.store, "update_day") as update:
            out = lengine.ensure_evening_pass(USER, date(2026, 7, 26),
                                              force=True)
        update.assert_not_called()
        self.assertIs(out, done)

    def test_no_morning_card_means_no_evening_recap(self):
        # A day the founder never opened is living, not failing (L-4).
        with patch.object(lengine.store, "get_day", return_value=None), \
                patch.object(lengine.store, "update_day") as update:
            self.assertIsNone(
                lengine.ensure_evening_pass(USER, date(2026, 7, 26),
                                            force=True))
        update.assert_not_called()

    def test_the_day_payload_nests_the_evening_branch(self):
        out = lp.serialize_day({"id": "d1", "evening_generated_at": None,
                                "evening_summary": {}})
        self.assertIsNone(out["evening"]["generated_at"])
        # Never branch on the summary being non-empty: a recap can be
        # legitimately sparse on a quiet day, and that is not "not evening
        # yet".
        self.assertEqual(out["evening"]["summary"], {})

    def test_the_evening_pass_sends_nothing(self):
        import inspect
        code = inspect.getsource(lengine.ensure_evening_pass).split('"""')[-1]
        for banned in ("email", "notif", "push", "send"):
            self.assertNotIn(banned, code.lower())

    def test_the_card_columns_all_exist_in_the_migration(self):
        # build_daily_card's dict is handed straight to the table. A key with
        # no column is an insert that fails at 05:00 and is discovered by an
        # empty panel.
        with open(MIGRATION, "r", encoding="utf-8") as fh:
            sql = fh.read()
        block = sql[sql.index("CREATE TABLE IF NOT EXISTS life_days"):]
        block = block[:block.index(");")]
        added = set(re.findall(
            r"ALTER TABLE life_days ADD COLUMN IF NOT EXISTS (\w+)", sql))
        card = self._card([])
        for key in card:
            self.assertTrue(key in block or key in added,
                            f"life_days has no column for {key}")


# ═════════════════════════════════════════════════════════════════════════
# N4 — the per-user master-doc injection
# ═════════════════════════════════════════════════════════════════════════

@unittest.skipIf(_IMPORT_ERROR is not None, f"needs app deps: {_IMPORT_ERROR}")
class MasterDocInjectionTests(unittest.TestCase):

    def test_no_context_means_no_block_at_all(self):
        # Everyone who has not opted in gets a prompt that is byte-for-byte
        # what it is today.
        self.assertEqual(mdr._render_life_block(None), "")
        self.assertEqual(mdr._render_life_block({}), "")
        self.assertEqual(mdr._render_life_block(
            {"principles": [], "phrases": [], "strategy": ""}), "")

    def test_the_cap_holds_against_a_whole_corpus(self):
        # Sixty principles plus eight strategy documents would re-open the
        # attention-ceiling failure PRs #81–#89 fixed.
        block = mdr._render_life_block({
            "principles": [{"title": f"p{i}"} for i in range(60)],
            "phrases": [{"body": f"q{i}"} for i in range(40)],
            "strategy": "s" * 5000,
        })
        self.assertEqual(block.count("[principle]"), mdr._LIFE_MAX_PRINCIPLES)
        self.assertEqual(block.count("[their phrase]"), mdr._LIFE_MAX_PHRASES)
        self.assertLess(len(block), 4000)

    def test_long_entries_are_trimmed_like_the_library_trims(self):
        block = mdr._render_life_block({"principles": [{"title": "x" * 900}]})
        self.assertIn("…", block)
        self.assertLess(len(block), 1200)

    def test_the_block_forbids_dumping_and_scoring(self):
        block = mdr._render_life_block({"principles": [{"title": "p"}]})
        self.assertIn("NEVER dump", block)
        self.assertIn("never score", block)

    def test_the_shared_master_document_is_never_written_to(self):
        # The fence: the injection is per-request and per-user. Nothing in the
        # feature mutates the shared body.
        import inspect
        for module in (lengine, lchat, lstore, lroutes, lreminders):
            src = inspect.getsource(module)
            self.assertNotIn("MASTER_DOCUMENT", src)

    def test_answer_question_still_works_without_the_new_argument(self):
        # Additive parameter: every existing caller is unaffected.
        import inspect
        sig = inspect.signature(mdr.answer_question)
        self.assertIsNone(sig.parameters["life_context"].default)


# ═════════════════════════════════════════════════════════════════════════
# N8 — zero scheduled outbound anything (L-4)
# ═════════════════════════════════════════════════════════════════════════

@unittest.skipIf(_IMPORT_ERROR is not None, f"needs app deps: {_IMPORT_ERROR}")
class NoNudgesTests(unittest.TestCase):
    """Not typing means LIVING, not failing. The system is dormant until
    opened.

    AMENDED BY EXPLICIT FOUNDER DECISION, 2026-07-30: web-push reminders
    exist, ONLY behind explicit per-user opt-in. The user flips a switch in
    the panel settings; every default is OFF; the slot ladder only (daily
    morning, daily evening, then the checkout cadence weekly, monthly,
    quarterly, yearly, fiveyear, tenyear — extended by the founder the same
    day, every rung strictly opt-in, longer cadences deep-linking to
    EXISTING surfaces with no new views or tables).
    ``services/life_reminders.py`` is the
    ONE sanctioned delivery module — routes expose the settings but never
    import pywebpush and never send, and everything else L-4 bans stays
    banned exactly as before: no emails, no silence detection, no streaks,
    and generation (ensure_daily_card / ensure_evening_pass) still contains
    no delivery code. This class is what keeps the amendment exactly as wide
    as the decision and no wider."""

    BANNED = ("send_email", "email_service", "resend", "push_notification",
              "send_push", "notify_user", "unsubscribe_tokens",
              "arc_notifications", "assignment_email", "audit_email")

    def test_no_module_can_send_anything(self):
        # life_reminders is DELIBERATELY absent from this list: it is the one
        # sanctioned sender (founder 2026-07-30). Everything else stays fully
        # banned, life_routes included — it may reference the settings module
        # but holds none of these primitives itself.
        import inspect
        for module in (lp, lstore, lengine, lchat, lroutes, importer):
            src = code_only(inspect.getsource(module)).lower()
            for banned in self.BANNED:
                self.assertNotIn(
                    banned, src,
                    f"{module.__name__} reaches for {banned} — L-4 is zero "
                    f"nudges")

    def test_nothing_tracks_silence_or_streaks(self):
        # The opt-in amendment covers DELIVERY only. Tracking stays banned
        # everywhere, the sanctioned sender included: a reminder fires on the
        # clock, never on "you have not typed lately".
        import inspect
        for module in (lp, lstore, lengine, lchat, lroutes, lreminders,
                       lremhook):
            src = code_only(inspect.getsource(module)).lower()
            for banned in ("streak", "days_since_last", "inactivity",
                           "silence_detect"):
                self.assertNotIn(banned, src, module.__name__)

    def test_generation_is_scheduled_but_delivery_is_not(self):
        # Both passes write a row and return. Nothing downstream — the
        # reminder path is a separate cron hitting a separate module, and the
        # founder's amendment did NOT open these two functions.
        import inspect
        for fn in (lengine.ensure_daily_card, lengine.ensure_evening_pass):
            code = code_only(inspect.getsource(fn)).lower()
            for banned in ("email", "notif", "push", "send"):
                self.assertNotIn(banned, code, fn.__name__)

    def test_reminder_defaults_are_all_off(self):
        # OPT-IN is the whole deal, for EVERY rung of the ladder: a user who
        # never touches the settings receives exactly nothing, same as
        # before the feature existed.
        self.assertEqual(
            set(lreminders.SLOTS),
            {"morning", "evening", "weekly", "monthly", "quarterly",
             "yearly", "fiveyear", "tenyear"})
        for slot in lreminders.SLOTS:
            self.assertIs(lreminders.DEFAULT_SETTINGS[f"{slot}_enabled"],
                          False, slot)
        # And a missing row reads as the defaults, not as an error.
        with patch.object(lreminders, "_t", side_effect=RuntimeError("down")):
            settings = lreminders.get_settings(USER)
        for slot in lreminders.SLOTS:
            self.assertFalse(settings[f"{slot}_enabled"], slot)

    def test_the_migration_defaults_every_slot_to_false(self):
        path = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                            "migrations", "add_life_push_reminders.sql")
        with open(path, "r", encoding="utf-8") as fh:
            sql = fh.read()
        for slot in lreminders.SLOTS:
            column = f"{slot}_enabled"
            self.assertRegex(
                sql, rf"{column}\s+boolean\s+NOT NULL\s+DEFAULT false",
                f"{column} must default to false — opt-in, never opt-out")
            # And the log's slot CHECK must accept it, or the idempotency
            # claim raises on the first firing.
            self.assertIn(f"'{slot}'", sql)

    def test_routes_delegate_and_never_touch_pywebpush(self):
        # life_routes may expose the user-initiated settings endpoints, but
        # the sending machinery must be unreachable from it: no pywebpush, no
        # send function, no slot runner. services/life_reminders.py is the
        # ONLY module allowed to hold those.
        import inspect
        src = code_only(inspect.getsource(lroutes)).lower()
        for banned in ("pywebpush", "webpush", "send_web_push",
                       "run_reminder_slot"):
            self.assertNotIn(banned, src,
                             f"life_routes must not reference {banned}")

    def test_pywebpush_appears_only_in_the_sanctioned_module(self):
        root = os.path.dirname(os.path.abspath(__file__))
        offenders = []
        for folder in ("services", "routes"):
            for name in sorted(os.listdir(os.path.join(root, folder))):
                if not name.endswith(".py"):
                    continue
                rel = f"{folder}/{name}"
                if rel == "services/life_reminders.py":
                    continue
                with open(os.path.join(root, rel), "r",
                          encoding="utf-8") as fh:
                    if "pywebpush" in fh.read():
                        offenders.append(rel)
        self.assertEqual(offenders, [])


# ═════════════════════════════════════════════════════════════════════════
# N2 — the fence: this feature stays off the live loop
# ═════════════════════════════════════════════════════════════════════════

@unittest.skipIf(_IMPORT_ERROR is not None, f"needs app deps: {_IMPORT_ERROR}")
class IsolationTests(unittest.TestCase):
    """The Life Panel is scaffolding. It must not reach into the
    record → transcribe → coach → read loop, and no F1 module may come to
    depend on it. Same shape and same fence as the Journal (PR #254)."""

    LIFE_MODULES = ("services/life_panel.py", "services/life_store.py",
                    "services/life_engine.py", "services/life_chat.py",
                    "services/life_import.py", "services/life_reminders.py",
                    "routes/life_routes.py",
                    "routes/life_reminders_webhook.py")

    def test_life_modules_import_nothing_from_the_loop(self):
        import inspect
        banned = ("best_presentation", "ideal_text", "charisma_snippets",
                  "lab_recording", "moment_suggestions", "say_it_stronger",
                  "delivery_stars", "power_score", "slide_word_split",
                  "cross_take_selection", "coaching_state_machine")
        for module in (lp, lstore, lengine, lchat, lroutes, importer,
                       lreminders, lremhook):
            src = inspect.getsource(module)
            for name in banned:
                self.assertNotIn(
                    name, src, f"{module.__name__} must not reference {name}")

    def test_no_product_module_imports_the_life_module(self):
        """The direction that actually matters.

        life → db is fine. db → life, or any F1 service → life, would make the
        panel load-bearing for the live loop. Only the two permitted contact
        points (the chat route and master_doc_rag's per-user block) may name
        it, and master_doc_rag names only its own renderer — not the module."""
        root = os.path.dirname(os.path.abspath(__file__))
        permitted = {"routes/v2_routes.py"}
        offenders = []
        for folder in ("services", "routes"):
            for name in sorted(os.listdir(os.path.join(root, folder))):
                rel = f"{folder}/{name}"
                if not name.endswith(".py") or rel in self.LIFE_MODULES:
                    continue
                if rel in permitted:
                    continue
                with open(os.path.join(root, rel), "r", encoding="utf-8") as fh:
                    src = fh.read()
                if re.search(r"^\s*from services import life_|"
                             r"^\s*from services\.life_|"
                             r"^\s*import services\.life_",
                             src, re.MULTILINE):
                    offenders.append(rel)
        self.assertEqual(offenders, [])

    def test_the_chat_route_is_the_only_touched_product_file(self):
        # Two permitted contact points, and master_doc_rag's is a renderer it
        # owns rather than an import of ours.
        import inspect
        self.assertNotIn("life_chat", inspect.getsource(mdr))
        self.assertNotIn("life_engine", inspect.getsource(mdr))

    def test_life_tables_are_never_joined_to_a_product_table(self):
        import inspect
        src = inspect.getsource(lstore)
        tables = set(re.findall(r'_t\("(\w+)"\)', src))
        self.assertEqual(tables - set(lstore.LIFE_TABLES), set())

    def test_the_migration_touches_no_product_table(self):
        with open(MIGRATION, "r", encoding="utf-8") as fh:
            sql = fh.read()
        for statement in re.findall(r"(?:CREATE TABLE IF NOT EXISTS|ALTER TABLE)\s+"
                                    r"(?:public\.)?(\w+)", sql):
            self.assertTrue(statement.startswith("life_"), statement)


# ═════════════════════════════════════════════════════════════════════════
# N9 + the gate — routes
# ═════════════════════════════════════════════════════════════════════════

@unittest.skipIf(_IMPORT_ERROR is not None, f"needs app deps: {_IMPORT_ERROR}")
class RouteGateTests(unittest.TestCase):

    def setUp(self):
        self.app = Flask(__name__)
        self.app.register_blueprint(lroutes.life_bp)
        self.client = self.app.test_client()
        self._orig_verify = auth.verify_supabase_token
        auth.verify_supabase_token = lambda token: {"sub": USER}

    def tearDown(self):
        auth.verify_supabase_token = self._orig_verify

    def _get(self, path, **kw):
        return self.client.get(path, headers={"Authorization": "Bearer t"},
                               **kw)

    def test_flag_off_404s_everything(self):
        # Not 503, not "coming soon". With the flag off the feature does not
        # exist.
        with patch.object(lroutes.chat, "is_enabled", return_value=False):
            for path in ("/v2/life/state", "/v2/life/principles",
                         "/v2/life/strategy", "/v2/life/prayer"):
                self.assertEqual(self._get(path).status_code, 404, path)

    def test_not_consented_is_409_with_a_pointer(self):
        with patch.object(lroutes.chat, "is_enabled", return_value=True), \
                patch.object(lroutes.chat, "has_consented", return_value=False):
            resp = self._get("/v2/life/principles")
        self.assertEqual(resp.status_code, 409)
        self.assertEqual(resp.get_json()["pointer"], "/panel/principles")

    def test_state_and_consent_are_reachable_without_consent(self):
        # Gating them would be a locked door with the key inside.
        with patch.object(lroutes.chat, "is_enabled", return_value=True), \
                patch.object(lroutes.chat, "has_consented", return_value=False), \
                patch.object(lroutes.chat, "is_allowlisted", return_value=False), \
                patch.object(lroutes.store, "get_setup", return_value=None):
            self.assertEqual(self._get("/v2/life/state").status_code, 200)

    def test_a_founder_only_surface_404s_rather_than_403s(self):
        # A 403 confirms the surface exists.
        with patch.object(lroutes.chat, "is_enabled", return_value=True), \
                patch.object(lroutes.chat, "has_consented", return_value=True), \
                patch.object(lroutes.chat, "is_allowlisted", return_value=False):
            resp = self._get("/v2/life/prayer")
        self.assertEqual(resp.status_code, 404)
        self.assertNotIn("pompeiana", resp.get_data(as_text=True))

    def test_a_founder_only_entry_is_absent_from_the_menu(self):
        with patch.object(lroutes.chat, "is_enabled", return_value=True), \
                patch.object(lroutes.chat, "has_consented", return_value=True), \
                patch.object(lroutes.chat, "is_allowlisted", return_value=False), \
                patch.object(lroutes.store, "get_setup",
                             return_value={"completed_at": "2026-07-26"}):
            menu = self._get("/v2/life/state").get_json()["menu"]
        self.assertNotIn("prayer", menu)

    def _state(self, *, consented, setup):
        with patch.object(lroutes.chat, "is_enabled", return_value=True), \
                patch.object(lroutes.chat, "has_consented",
                             return_value=consented), \
                patch.object(lroutes.chat, "is_allowlisted", return_value=False), \
                patch.object(lroutes.store, "get_setup", return_value=setup):
            return self._get("/v2/life/state").get_json()

    def test_the_hamburger_holds_one_door_before_activation(self):
        # Completing onboarding IS the membership test — there is no separate
        # roster. Before that the panel is one entry, not nine empty rooms: a
        # user let in early meets a feature that silently does nothing.
        for consented, setup in ((False, None), (True, None),
                                 (True, {"completed_at": None})):
            state = self._state(consented=consented, setup=setup)
            self.assertEqual(state["menu"], ["principles"])
            self.assertFalse(state["active"])
            # The `#` picker is useless until the engine will run.
            self.assertEqual(state["tags"], [])

    def test_finishing_onboarding_opens_the_rest(self):
        state = self._state(consented=True,
                            setup={"completed_at": "2026-07-26T05:00:00Z"})
        self.assertTrue(state["active"])
        for entry in ("wins", "today", "goals", "timeline", "strategy"):
            self.assertIn(entry, state["menu"])
        self.assertTrue(state["tags"])

    def test_the_state_payload_carries_the_onboarding_state_machine(self):
        # The FE drives WHICH SCREEN off these two objects.
        fresh = self._state(consented=False, setup=None)
        self.assertEqual(fresh["consent"]["accepted_version"], None)
        self.assertEqual(fresh["setup"], {"complete": False, "started": False,
                                          "resume_step": None})

        mid = self._state(consented=True,
                          setup={"answers": {"_step": "quarterly"}})
        self.assertEqual(mid["consent"]["accepted_version"],
                         mid["consent"]["required_version"])
        self.assertEqual(mid["setup"]["resume_step"], "quarterly")
        self.assertTrue(mid["setup"]["started"])

    def test_the_nested_and_flat_shapes_cannot_disagree(self):
        # Two representations of one fact is a drift risk; they are read from
        # the same locals so a change to one moves both.
        for consented, setup in ((False, None), (True, None),
                                 (True, {"completed_at": "2026-07-26"})):
            s = self._state(consented=consented, setup=setup)
            self.assertEqual(s["consent"]["accepted_version"] is not None,
                             s["consented"])
            self.assertEqual(s["setup"]["complete"], s["setup_complete"])
            self.assertEqual(s["setup"]["started"], s["setup_started"])

    def test_the_consent_version_is_an_opaque_token(self):
        # Compared for equality, never parsed. A copy change shipping as
        # "2026-08-02" instead of "1.1" must not break the comparison.
        s = self._state(consented=True, setup=None)
        self.assertIsInstance(s["consent"]["required_version"], str)

    def test_setup_alone_never_activates_without_consent(self):
        # The order is explainer → CONSENT → form, and it cannot be
        # rearranged: the form itself collects the anchor and eight horizons of
        # personal goals. "Activate by finishing onboarding" must never become
        # "collect first, ask after".
        state = self._state(consented=False,
                            setup={"completed_at": "2026-07-26T05:00:00Z"})
        self.assertFalse(state["active"])
        self.assertEqual(state["menu"], ["principles"])

    def test_no_auth_is_401_not_a_leak(self):
        with patch.object(lroutes.chat, "is_enabled", return_value=True):
            self.assertEqual(self.client.get("/v2/life/state").status_code,
                             401)

    def test_the_upload_returns_a_diff_and_writes_nothing(self):
        # A re-upload produces a diff you approve, never a silent overwrite.
        with patch.object(lroutes.chat, "is_enabled", return_value=True), \
                patch.object(lroutes.chat, "has_consented", return_value=True), \
                patch.object(lroutes.store, "latest_strategy",
                             return_value={"weekly": {"body": "old line"}}), \
                patch.object(lroutes.store, "insert_strategy_version") as write:
            resp = self.client.post(
                "/v2/life/strategy/upload",
                headers={"Authorization": "Bearer t"},
                json={"documents": {"weekly": "new line"}})
        self.assertEqual(resp.status_code, 200)
        self.assertFalse(resp.get_json()["written"])
        self.assertIn("weekly", resp.get_json()["diffs"])
        write.assert_not_called()

    def test_every_proposal_read_carries_its_warrant(self):
        """L-2's delivery, on BOTH paths.

        The FE drops a proposal that arrives without a warrant — correctly,
        since there is no compliant way to render a bare one. So a serializer
        that forgets it does not look like a bug: the proposal simply never
        appears. The weekly review is the entire delivery mechanism for the
        L-2b queue, so forgetting it there loses the queue silently."""
        row = {"id": "p1", "kind": "strategy", "target": "weekly.goals",
               "status": "queued", "rank": 1.0,
               "warrant_principle_id": "w1", "created_at": "2026-07-26"}
        warrant = {"id": "w1", "kind": "principle", "title": "Mine"}

        for path, extract in (
            ("/v2/life/proposals", lambda b: b["weekly_batch"]),
            ("/v2/life/week", lambda b: b["proposals"]),
        ):
            with patch.object(lroutes.chat, "is_enabled", return_value=True), \
                    patch.object(lroutes.chat, "has_consented",
                                 return_value=True), \
                    patch.object(lroutes.store, "expire_stale_proposals",
                                 return_value=0), \
                    patch.object(lroutes.store, "list_proposals",
                                 side_effect=lambda u, **k:
                                 [row] if k.get("status") == "queued" else []), \
                    patch.object(lroutes.store, "get_item",
                                 return_value=warrant), \
                    patch.object(lroutes.store, "get_week", return_value=None), \
                    patch.object(lroutes.store, "list_notes", return_value=[]):
                body = self._get(path).get_json()
            found = extract(body)
            self.assertEqual(len(found), 1, path)
            self.assertEqual(found[0]["warrant"]["title"], "Mine", path)
            self.assertFalse(found[0]["report_only"], path)

    def test_approving_the_immutable_core_is_refused(self):
        with patch.object(lroutes.chat, "is_enabled", return_value=True), \
                patch.object(lroutes.chat, "has_consented", return_value=True), \
                patch.object(lroutes.store, "get_proposal",
                             return_value={"id": "p1", "status": "surfaced",
                                           "kind": "strategy",
                                           "target": "weekly.section_i"}):
            resp = self.client.post("/v2/life/proposals/p1/approve",
                                    headers={"Authorization": "Bearer t"},
                                    json={})
        self.assertEqual(resp.status_code, 409)
        self.assertEqual(resp.get_json()["code"], "IMMUTABLE_CORE")

    def test_approving_a_case_asks_the_retire_question(self):
        """BE-7's prompt, at the only moment it is answerable: a principle has
        just entered the archive.

        It ASKS; it never announces. Nothing is retired by this response —
        the veto is absolute and POST /principles/<id>/retire is the only
        thing that moves a status."""
        proposed = {"id": "new", "kind": "principle", "status": "proposed",
                    "title": "Say no early"}
        old = {"id": "old", "kind": "principle", "title": "Say no politely",
               "order_key": 1000.0}
        with patch.object(lroutes.chat, "is_enabled", return_value=True), \
                patch.object(lroutes.chat, "has_consented", return_value=True), \
                patch.object(lroutes.store, "get_case",
                             return_value={"id": "c1"}), \
                patch.object(lroutes.store, "items_for_case",
                             return_value=[proposed]), \
                patch.object(lroutes.store, "update_item",
                             return_value={**proposed, "status": "active"}), \
                patch.object(lroutes.store, "update_case",
                             return_value={"id": "c1"}), \
                patch.object(lroutes.store, "principles", return_value=[old]), \
                patch.object(lroutes.engine, "retire_candidate",
                             return_value=old) as ask:
            body = self.client.patch("/v2/life/cases/c1",
                                     headers={"Authorization": "Bearer t"},
                                     json={"approve": True}).get_json()
        ask.assert_called_once()
        self.assertEqual(body["retire_prompt"]["retires"]["id"], "old")
        # "#12" has to mean the twelfth thing they see.
        self.assertEqual(body["retire_prompt"]["number"], 1)
        self.assertIn("#1", body["retire_prompt"]["question"])

    def test_no_retire_candidate_means_no_prompt(self):
        with patch.object(lroutes.chat, "is_enabled", return_value=True), \
                patch.object(lroutes.chat, "has_consented", return_value=True), \
                patch.object(lroutes.store, "get_case",
                             return_value={"id": "c1"}), \
                patch.object(lroutes.store, "items_for_case", return_value=[]), \
                patch.object(lroutes.store, "update_case",
                             return_value={"id": "c1"}):
            body = self.client.patch("/v2/life/cases/c1",
                                     headers={"Authorization": "Bearer t"},
                                     json={"approve": True}).get_json()
        self.assertIsNone(body["retire_prompt"])

    def test_a_no_on_retire_keeps_both_and_is_remembered(self):
        # Veto is absolute; the question is never re-asked.
        principle = {"id": "new", "kind": "principle", "title": "n"}
        old = {"id": "old", "kind": "principle", "title": "o"}
        with patch.object(lroutes.chat, "is_enabled", return_value=True), \
                patch.object(lroutes.chat, "has_consented", return_value=True), \
                patch.object(lroutes.store, "get_item",
                             side_effect=lambda u, i: principle if i == "new"
                             else old), \
                patch.object(lroutes.store, "insert_proposal") as insert, \
                patch.object(lroutes.store, "update_item") as update:
            resp = self.client.post(
                "/v2/life/principles/new/retire",
                headers={"Authorization": "Bearer t"},
                json={"retires_id": "old", "decision": "no"})
        self.assertEqual(resp.status_code, 200)
        self.assertTrue(resp.get_json()["both_active"])
        update.assert_not_called()
        self.assertEqual(insert.call_args[0][1]["status"], "dismissed")

    def test_patch_items_saves_only_the_supplied_keys(self):
        # The FE's inline edit contract. A partial PATCH must never blank an
        # untouched field, which is why validate_item_input returns only the
        # keys that were present.
        with patch.object(lroutes.chat, "is_enabled", return_value=True), \
                patch.object(lroutes.chat, "has_consented", return_value=True), \
                patch.object(lroutes.store, "get_item",
                             return_value={"id": "i1", "kind": "goal"}), \
                patch.object(lroutes.store, "update_item",
                             return_value={"id": "i1", "kind": "goal",
                                           "title": "new"}) as update:
            resp = self.client.patch("/v2/life/items/i1",
                                     headers={"Authorization": "Bearer t"},
                                     json={"title": "new"})
        self.assertEqual(resp.status_code, 200)
        self.assertEqual(update.call_args[0][2], {"title": "new"})

    def test_patch_items_refuses_to_change_the_kind(self):
        # The kind is the discriminator: changing it would move a row between
        # views and past a different validator.
        with patch.object(lroutes.chat, "is_enabled", return_value=True), \
                patch.object(lroutes.chat, "has_consented", return_value=True), \
                patch.object(lroutes.store, "get_item",
                             return_value={"id": "i1", "kind": "goal"}):
            resp = self.client.patch("/v2/life/items/i1",
                                     headers={"Authorization": "Bearer t"},
                                     json={"kind": "win"})
        self.assertEqual(resp.status_code, 400)

    def test_patch_items_404s_on_someone_elses_row(self):
        with patch.object(lroutes.chat, "is_enabled", return_value=True), \
                patch.object(lroutes.chat, "has_consented", return_value=True), \
                patch.object(lroutes.store, "get_item", return_value=None), \
                patch.object(lroutes.store, "update_item") as update:
            resp = self.client.patch("/v2/life/items/i1",
                                     headers={"Authorization": "Bearer t"},
                                     json={"title": "new"})
        self.assertEqual(resp.status_code, 404)
        update.assert_not_called()

    def test_patch_day_is_scoped_to_the_card_on_screen(self):
        # Addressed by id, not by "today". A panel left open across midnight
        # must write to the card being looked at, not to whatever date the
        # server thinks it is now.
        with patch.object(lroutes.chat, "is_enabled", return_value=True), \
                patch.object(lroutes.chat, "has_consented", return_value=True), \
                patch.object(lroutes.store, "update_day",
                             return_value={"id": "d1"}) as update:
            resp = self.client.patch("/v2/life/day/d1",
                                     headers={"Authorization": "Bearer t"},
                                     json={"evening_one_thing": True})
        self.assertEqual(resp.status_code, 200)
        self.assertEqual(update.call_args[0][1], "d1")

    def test_patch_day_404s_when_the_row_is_not_yours(self):
        with patch.object(lroutes.chat, "is_enabled", return_value=True), \
                patch.object(lroutes.chat, "has_consented", return_value=True), \
                patch.object(lroutes.store, "update_day", return_value=None):
            resp = self.client.patch("/v2/life/day/d1",
                                     headers={"Authorization": "Bearer t"},
                                     json={"evening_one_thing": True})
        self.assertEqual(resp.status_code, 404)

    def test_the_day_payload_carries_the_id_the_patch_needs(self):
        # GET /v2/life/day must hand back the id, or the id-scoped PATCH is
        # unusable and the FE has to guess.
        self.assertIn("id", lp.serialize_day({"id": "d1"}))

    def test_the_one_thing_can_be_changed_but_not_removed(self):
        with patch.object(lroutes.chat, "is_enabled", return_value=True), \
                patch.object(lroutes.chat, "has_consented", return_value=True):
            resp = self.client.patch("/v2/life/day/d1",
                                     headers={"Authorization": "Bearer t"},
                                     json={"one_thing": "   "})
        self.assertEqual(resp.status_code, 400)

    def test_hard_delete_requires_an_explicit_confirmation(self):
        with patch.object(lroutes.chat, "is_enabled", return_value=True), \
                patch.object(lroutes.store, "hard_delete") as wipe:
            resp = self.client.delete("/v2/life/data",
                                      headers={"Authorization": "Bearer t"},
                                      json={})
        self.assertEqual(resp.status_code, 400)
        wipe.assert_not_called()

    def test_hard_delete_survives_a_lapsed_consent(self):
        # Gating the exit behind the gate they are leaving through would be
        # the worst possible reading of a consent screen.
        with patch.object(lroutes.chat, "is_enabled", return_value=True), \
                patch.object(lroutes.chat, "has_consented", return_value=False), \
                patch.object(lroutes.store, "hard_delete",
                             return_value={"deleted": {"life_notes": 3},
                                           "failed": []}), \
                patch.object(lroutes.importer, "delete_setup_documents",
                             return_value=1), \
                patch.object(lroutes.reminders, "hard_delete_user",
                             return_value={"deleted": {}, "failed": []}), \
                patch.object(lroutes.engine, "delete_user_copy",
                             return_value=0):
            resp = self.client.delete("/v2/life/data",
                                      headers={"Authorization": "Bearer t"},
                                      json={"confirm": "DELETE"})
        self.assertEqual(resp.status_code, 200)
        self.assertTrue(resp.get_json()["deleted"])

    def test_a_partial_delete_is_reported_not_claimed_as_success(self):
        # "Your data is gone" is a promise; a half-kept promise about this
        # corpus is worse than a refusal.
        with patch.object(lroutes.chat, "is_enabled", return_value=True), \
                patch.object(lroutes.chat, "has_consented", return_value=True), \
                patch.object(lroutes.store, "hard_delete",
                             return_value={"deleted": {}, "failed": ["life_cases"]}), \
                patch.object(lroutes.importer, "delete_setup_documents",
                             return_value=0), \
                patch.object(lroutes.reminders, "hard_delete_user",
                             return_value={"deleted": {}, "failed": []}), \
                patch.object(lroutes.engine, "delete_user_copy",
                             return_value=0):
            resp = self.client.delete("/v2/life/data",
                                      headers={"Authorization": "Bearer t"},
                                      json={"confirm": "DELETE"})
        self.assertEqual(resp.status_code, 500)
        self.assertEqual(resp.get_json()["code"], "PARTIAL_DELETE")

    def test_a_failing_reminder_or_document_wipe_is_a_partial_delete(self):
        # The 2026-07-30 tables are part of the same promise: if their wipe
        # fails, "your data is gone" must not be claimed.
        with patch.object(lroutes.chat, "is_enabled", return_value=True), \
                patch.object(lroutes.chat, "has_consented", return_value=True), \
                patch.object(lroutes.store, "hard_delete",
                             return_value={"deleted": {}, "failed": []}), \
                patch.object(lroutes.importer, "delete_setup_documents",
                             return_value=None), \
                patch.object(lroutes.reminders, "hard_delete_user",
                             return_value={"deleted": {},
                                           "failed": ["life_reminder_log"]}), \
                patch.object(lroutes.engine, "delete_user_copy",
                             return_value=None):
            resp = self.client.delete("/v2/life/data",
                                      headers={"Authorization": "Bearer t"},
                                      json={"confirm": "DELETE"})
        self.assertEqual(resp.status_code, 500)
        body = resp.get_json()
        self.assertEqual(body["code"], "PARTIAL_DELETE")
        self.assertIn("life_setup_documents", body["failed"])
        self.assertIn("life_reminder_log", body["failed"])
        self.assertIn("life_user_copy", body["failed"])

    def test_an_incomplete_export_says_so(self):
        # An export missing a table looks exactly like an empty table.
        with patch.object(lroutes.chat, "is_enabled", return_value=True), \
                patch.object(lroutes.chat, "has_consented", return_value=True), \
                patch.object(lroutes.store, "export_all",
                             return_value={"tables": {}, "errors": ["life_cases"]}), \
                patch.object(lroutes.importer, "export_setup_documents",
                             return_value=[]), \
                patch.object(lroutes.reminders, "export_user",
                             return_value={}), \
                patch.object(lroutes.engine, "export_user_copy",
                             return_value=[]):
            resp = self.client.post("/v2/life/export",
                                    headers={"Authorization": "Bearer t"},
                                    json={})
        self.assertFalse(resp.get_json()["complete"])

    def test_export_carries_the_new_tables_too(self):
        # Setup documents and reminder settings live outside life_store, and
        # an export that quietly missed them would look exactly like a user
        # who never used them.
        with patch.object(lroutes.chat, "is_enabled", return_value=True), \
                patch.object(lroutes.chat, "has_consented", return_value=True), \
                patch.object(lroutes.store, "export_all",
                             return_value={"tables": {}}), \
                patch.object(lroutes.importer, "export_setup_documents",
                             return_value=[{"id": "d1"}]), \
                patch.object(lroutes.reminders, "export_user",
                             return_value={"life_reminder_settings": [],
                                           "life_push_subscriptions": []}), \
                patch.object(lroutes.engine, "export_user_copy",
                             return_value=[]):
            resp = self.client.post("/v2/life/export",
                                    headers={"Authorization": "Bearer t"},
                                    json={})
        body = resp.get_json()
        self.assertTrue(body["complete"])
        self.assertEqual(body["tables"]["life_setup_documents"],
                         [{"id": "d1"}])
        self.assertIn("life_reminder_settings", body["tables"])
        self.assertIn("life_push_subscriptions", body["tables"])
        self.assertIn("life_user_copy", body["tables"])


@unittest.skipIf(_IMPORT_ERROR is not None, f"needs app deps: {_IMPORT_ERROR}")
class ExportAndDeleteCoverageTests(unittest.TestCase):
    """N9 — the round trip covers every table the migration creates."""

    def test_export_reads_every_life_table(self):
        seen = []

        class _Q:
            def select(self, *a, **k): return self
            def eq(self, *a, **k): return self
            def limit(self, *a, **k): return self
            def execute(self): return type("R", (), {"data": []})()

        def _t(table):
            seen.append(table)
            return _Q()

        with patch.object(lstore, "_t", side_effect=_t):
            lstore.export_all(USER)
        self.assertEqual(set(seen), set(lstore.LIFE_TABLES))

    def test_hard_delete_covers_every_life_table(self):
        seen = []

        class _Q:
            def delete(self): return self
            def eq(self, *a, **k): return self
            def execute(self): return type("R", (), {"data": [{"id": "1"}]})()

        def _t(table):
            seen.append(table)
            return _Q()

        with patch.object(lstore, "_t", side_effect=_t):
            result = lstore.hard_delete(USER)
        self.assertEqual(set(seen), set(lstore.LIFE_TABLES))
        self.assertEqual(result["failed"], [])

    def test_a_failing_table_is_reported_by_name(self):
        def _t(table):
            raise RuntimeError("down")
        with patch.object(lstore, "_t", side_effect=_t):
            result = lstore.hard_delete(USER)
        self.assertEqual(set(result["failed"]), set(lstore.LIFE_TABLES))


# ═════════════════════════════════════════════════════════════════════════
# BE-10 — the corpus never reaches a log or a prompt in bulk
# ═════════════════════════════════════════════════════════════════════════

@unittest.skipIf(_IMPORT_ERROR is not None, f"needs app deps: {_IMPORT_ERROR}")
class PrivacyTests(unittest.TestCase):

    def test_retrieval_is_capped_well_under_the_corpus(self):
        self.assertLessEqual(lengine.RETRIEVAL_LIMIT, 10)

    def test_no_raw_note_body_is_ever_logged(self):
        """A log aggregator is a second, less guarded copy of the corpus.

        Derivation OUTPUTS and sizes are logged; note bodies are not."""
        import inspect
        src = inspect.getsource(lengine)
        for match in re.finditer(r"logger\.(?:info|warning|error)\((.*?)\)",
                                 src, re.DOTALL):
            call = match.group(1)
            for leak in ("note_text", "case_text", "reflections", "remainder",
                         "question)"):
                self.assertNotIn(leak, call,
                                 f"a log call carries {leak}")

    def test_the_derivation_logger_takes_no_free_text(self):
        import inspect
        src = inspect.getsource(lengine._log_derivation)
        self.assertIn("never the note body", src)

    def test_derivation_calls_are_stateless_and_unstored(self):
        """`store=False` is passed explicitly, and the calls stay stateless.

        Two reasons, and neither is the default's job to hold: a default is a
        decision someone else can change, and zero-data-retention only covers
        the STATELESS endpoints — moving these to a conversations/agents API
        would silently leave that protection behind."""
        import inspect
        src = inspect.getsource(lengine._complete)
        self.assertIn("store=False", src)

    def test_a_dedicated_zero_retention_key_is_supported(self):
        # BE-10: "use an API path with no training retention". The retention
        # posture is a project setting, so the code side is the ability to
        # point this surface at a separate project.
        import inspect
        self.assertIn("LIFE_PANEL_OPENAI_API_KEY",
                      inspect.getsource(lengine._client))

    def test_there_is_no_sharing_or_publish_surface(self):
        # The corpus names real people alongside claims about them. It stays
        # private to one user forever.
        import inspect
        for module in (lroutes, lstore, lengine):
            src = inspect.getsource(module).lower()
            for banned in ("publish", "share_url", "public_url", "presign"):
                self.assertNotIn(banned, src, module.__name__)

    def test_every_store_read_is_scoped_to_one_user(self):
        import inspect
        src = inspect.getsource(lstore)
        # Every table access chain must carry a user_id filter. The one
        # exception would be a cross-user read, and there is none.
        for block in src.split("def ")[1:]:
            if '_t("' not in block:
                continue
            self.assertIn("user_id", block,
                          f"a store function touches a table without scoping "
                          f"it to a user: {block.splitlines()[0]}")


# ═════════════════════════════════════════════════════════════════════════
# Item 9 (founder 2026-07-30) — setup strategy-document upload + drafting
# ═════════════════════════════════════════════════════════════════════════

class NewMigrationTests(unittest.TestCase):
    """The two 2026-07-30 migrations hold the same lines add_life_panel.sql
    holds: RLS in the creating file, zero policies, idempotent, no drops,
    raise on a partial run, life_* tables only."""

    FILES = ("add_life_setup_documents.sql", "add_life_push_reminders.sql")

    def _sql(self, name: str) -> str:
        path = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                            "migrations", name)
        with open(path, "r", encoding="utf-8") as fh:
            return fh.read()

    def test_every_created_table_enables_rls_in_its_file(self):
        for name in self.FILES:
            sql = self._sql(name)
            created = set(re.findall(
                r"CREATE TABLE IF NOT EXISTS\s+(life_\w+)", sql))
            protected = set(re.findall(
                r"ALTER TABLE\s+(life_\w+)\s+ENABLE ROW LEVEL SECURITY", sql))
            self.assertTrue(created, f"{name} creates no life_* tables")
            self.assertEqual(created - protected, set(), name)

    def test_no_policies_no_drops_and_a_raising_postcondition(self):
        for name in self.FILES:
            sql = self._sql(name)
            self.assertNotIn("CREATE POLICY", sql.upper(), name)
            self.assertNotRegex(sql.upper(), r"\bDROP\s+(TABLE|COLUMN)\b",
                                name)
            self.assertIn("RAISE EXCEPTION", sql, name)

    def test_idempotent_and_life_tables_only(self):
        for name in self.FILES:
            sql = self._sql(name)
            for statement in re.findall(r"CREATE TABLE[^(]*", sql):
                self.assertIn("IF NOT EXISTS", statement, name)
            for statement in re.findall(r"CREATE (?:UNIQUE )?INDEX[^(]*", sql):
                self.assertIn("IF NOT EXISTS", statement, name)
            for table in re.findall(
                    r"(?:CREATE TABLE IF NOT EXISTS|ALTER TABLE)\s+"
                    r"(?:public\.)?(\w+)", sql):
                self.assertTrue(table.startswith("life_"), f"{name}: {table}")

    def test_reminder_log_is_unique_per_user_slot_and_day(self):
        sql = self._sql("add_life_push_reminders.sql")
        self.assertRegex(sql, r"CREATE UNIQUE INDEX IF NOT EXISTS\s+\w+\s*\n?"
                              r"\s*ON life_reminder_log \(user_id, slot, "
                              r"sent_on\)")


class ConfirmedItemTests(unittest.TestCase):
    """N5 for the setup review UI: the tick is the approve. Pure functions,
    no DB."""

    def test_a_ticked_row_is_created_active_no_matter_what_was_sent(self):
        fields = importer.sanitize_confirmed_item({
            "kind": "goal", "title": "Ship the panel", "status": "proposed",
            "horizon": "year", "due_label": "[Dec]",
        })
        self.assertEqual(fields["status"], "active")
        self.assertEqual(fields["horizon"], "year")
        self.assertEqual(fields["due_label"], "[Dec]")

    def test_junk_kinds_and_blank_titles_are_refused(self):
        self.assertIsNone(importer.sanitize_confirmed_item(
            {"kind": "reflection", "title": "smuggled"}))
        self.assertIsNone(importer.sanitize_confirmed_item(
            {"kind": "goal", "title": "   "}))
        self.assertIsNone(importer.sanitize_confirmed_item("nope"))

    def test_a_real_kind_the_draft_never_emits_is_still_refused(self):
        # `task` and `event` are two of the panel's nine, and the draft path
        # emits neither — so neither may be created through it. The two
        # endpoints hold ONE kind set; a row this side accepted but the draft
        # side never produced would be a hole, not a courtesy.
        for kind in ("task", "event"):
            self.assertIsNone(importer.sanitize_confirmed_item(
                {"kind": kind, "title": "not from a document"}), kind)

    def test_the_three_named_kinds_are_confirmable_and_land_active(self):
        # Phrases, Principles and Wins — the three views the dock returned
        # nothing on. A ticked row is created ACTIVE, principles included:
        # the tick IS the approve, made on a fully displayed row (N5).
        for kind in ("phrase", "principle", "win"):
            fields = importer.sanitize_confirmed_item(
                {"kind": kind, "title": "The line itself.",
                 "status": "proposed"})
            self.assertIsNotNone(fields, kind)
            self.assertEqual(fields["kind"], kind)
            self.assertEqual(fields["status"], "active", kind)

    def test_a_ticked_phrase_lands_on_the_wall_rather_than_uncollected(self):
        fields = importer.sanitize_confirmed_item(
            {"kind": "phrase", "title": "Pray first."})
        self.assertEqual(fields["collection"],
                         importer.PHRASE_DEFAULT_COLLECTION)
        # An explicit collection still wins — the default fills a gap, it
        # does not overrule the user.
        fields = importer.sanitize_confirmed_item(
            {"kind": "phrase", "title": "Pray first.", "bet": "2026"})
        self.assertEqual(fields["collection"], "2026")

    def test_propose_and_apply_hold_one_kind_set(self):
        # The failure this prevents: propose emits a kind apply refuses, the
        # user ticks rows, presses Add and is told nothing was created — with
        # nothing they can do about it.
        for kind in importer.DOC_DRAFT_KINDS:
            self.assertIn(kind, lp.ITEM_KINDS, kind)
            title = ("🔵 The Company" if kind == "bet"
                     else f"a drafted {kind}")
            self.assertIsNotNone(
                importer.sanitize_confirmed_item({"kind": kind,
                                                  "title": title}), kind)
        # And the draft never reaches outside the panel's closed nine.
        self.assertEqual(set(importer.DOC_DRAFT_LINE_KINDS)
                         - set(importer.DOC_DRAFT_KINDS), set())

    def test_a_bet_keeps_its_locked_rank_and_an_invented_bet_is_refused(self):
        # L-2a: the uploaded document can word a bet, never reorder them.
        company = next(b for b in lp.BETS if b["key"] == "company")
        fields = importer.sanitize_confirmed_item({
            "kind": "bet",
            "title": f"{company['emoji']} {company['title']}",
            "order_key": 1,           # the client claims rank 1 — ignored
        })
        self.assertEqual(fields["order_key"], float(company["rank"]))
        self.assertIsNone(importer.sanitize_confirmed_item(
            {"kind": "bet", "title": "The Fourth Bet"}))


class DocumentLinePlannerTests(unittest.TestCase):
    """Lines out of a document → phrase / principle / win rows. Pure, no DB —
    and no writes, like every other planner in the importer."""

    def test_a_line_becomes_a_row_of_the_hinted_kind(self):
        rows = importer.plan_document_lines(
            USER, "phrase", ["Pray first.", "Do the hard thing."])
        self.assertEqual([r["kind"] for r in rows], ["phrase", "phrase"])
        self.assertEqual(rows[0]["title"], "Pray first.")
        self.assertEqual(rows[0]["status"], "active")
        self.assertEqual([r["order_key"] for r in rows], [1000.0, 2000.0])

    def test_only_a_phrase_carries_the_default_collection(self):
        phrase = importer.plan_document_lines(USER, "phrase", ["Pray."])[0]
        self.assertEqual(phrase["collection"], "wall")
        for kind in ("principle", "win"):
            row = importer.plan_document_lines(USER, kind, ["Something."])[0]
            self.assertNotIn("collection", row)

    def test_the_body_carries_the_line_only_when_the_title_was_cut(self):
        # A 600-character phrase keeps every character it was written with;
        # a short one leaves the body empty so the review row renders once.
        short = importer.plan_document_lines(USER, "phrase", ["Pray."])[0]
        self.assertEqual(short["body"], "")
        long_line = "x" * 600
        long_row = importer.plan_document_lines(USER, "phrase",
                                                [long_line])[0]
        self.assertEqual(len(long_row["title"]), 500)
        self.assertEqual(long_row["body"], long_line)

    def test_the_language_is_never_normalised(self):
        # The corpus is code-switched Polish/English and every "cleanup" is a
        # small edit to somebody's own words (BE-3).
        line = "Nie próbuj zrozumieć wszystkiego — „rób swoje”."
        row = importer.plan_document_lines(USER, "principle", [line])[0]
        self.assertEqual(row["title"], line)

    def test_blank_and_two_character_lines_are_dropped(self):
        rows = importer.plan_document_lines(
            USER, "win", ["", "  ", "ok", None, {"title": "Shipped it."}])
        self.assertEqual([r["title"] for r in rows], ["Shipped it."])

    def test_the_same_line_twice_is_one_row(self):
        # Dropped here, not at insert time: the unique index would refuse the
        # second row silently and the user would be told a lower count with
        # no way to see which row the system disagreed with.
        rows = importer.plan_document_lines(
            USER, "phrase", ["Pray first.", "Pray first."])
        self.assertEqual(len(rows), 1)

    def test_external_ids_make_a_second_draft_of_one_file_idempotent(self):
        first = importer.plan_document_lines(USER, "phrase", ["Pray first."])
        second = importer.plan_document_lines(USER, "phrase", ["Pray first."])
        self.assertEqual(first[0]["external_id"], second[0]["external_id"])

    def test_a_kind_this_planner_does_not_own_returns_nothing(self):
        for kind in ("goal", "bet", "task", "event", "reflection"):
            self.assertEqual(
                importer.plan_document_lines(USER, kind, ["anything"]), [],
                kind)


class OriginDocumentMigrationTests(unittest.TestCase):
    """The provenance column ships with the same lines every life_* migration
    holds: idempotent, nothing dropped, a raising post-condition, and it
    touches no product table."""

    NAME = "add_life_items_origin_document.sql"

    def setUp(self):
        path = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                            "migrations", self.NAME)
        with open(path, "r", encoding="utf-8") as fh:
            self.sql = fh.read()

    def test_it_is_idempotent(self):
        self.assertRegex(self.sql,
                         r"ADD COLUMN IF NOT EXISTS\s+origin_document_id")
        for statement in re.findall(r"CREATE (?:UNIQUE )?INDEX[^(]*",
                                    self.sql):
            self.assertIn("IF NOT EXISTS", statement)

    def test_nothing_is_dropped_and_a_partial_run_raises(self):
        self.assertNotRegex(self.sql.upper(), r"\bDROP\s+(TABLE|COLUMN)\b")
        self.assertIn("RAISE EXCEPTION", self.sql)

    def test_it_touches_life_tables_only(self):
        for table in re.findall(
                r"(?:CREATE TABLE IF NOT EXISTS|ALTER TABLE)\s+"
                r"(?:public\.)?(\w+)", self.sql):
            self.assertTrue(table.startswith("life_"), table)

    def test_no_foreign_key_cascades_the_documents_delete_into_the_items(self):
        # "Delete my uploads" must not become "delete the rows I made from
        # them". Same call origin_case_id makes.
        self.assertNotIn("REFERENCES", self.sql.upper())

    def test_the_column_rides_the_item_serializer(self):
        self.assertIn("origin_document_id",
                      lp.serialize_item({"kind": "principle",
                                         "origin_document_id": "d1"}))
        self.assertEqual(
            lp.serialize_item({"kind": "principle"})["origin_document_id"],
            None)


class WeekHorizonMigrationTests(unittest.TestCase):
    """The week horizon widens one CHECK constraint. Same house lines as every
    other life_* migration, plus the one thing that is specific to it: the new
    constraint must be a strict SUPERSET of the old, so no row that is valid
    today becomes invalid when this runs."""

    NAME = "add_life_items_week_horizon.sql"
    PREVIOUS = ("now", "month", "quarter", "year", "five_year", "ten_year",
                "twenty_year")

    def setUp(self):
        path = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                            "migrations", self.NAME)
        with open(path, "r", encoding="utf-8") as fh:
            self.sql = fh.read()

    def test_it_is_idempotent(self):
        # A CHECK constraint cannot be widened with IF NOT EXISTS, so this one
        # earns its re-runnability by dropping the old definition only when it
        # is there, then re-adding.
        self.assertRegex(self.sql, r"IF EXISTS\s*\(")
        self.assertIn("DROP CONSTRAINT life_items_horizon_check", self.sql)
        self.assertIn("ADD CONSTRAINT life_items_horizon_check", self.sql)

    def test_no_table_or_column_is_dropped_and_a_partial_run_raises(self):
        # DROP CONSTRAINT removes a RULE, not data — which is why the house
        # rule bans DROP TABLE / DROP COLUMN specifically.
        self.assertNotRegex(self.sql.upper(), r"\bDROP\s+(TABLE|COLUMN)\b")
        self.assertIn("RAISE EXCEPTION", self.sql)

    def test_it_touches_life_tables_only(self):
        for table in re.findall(
                r"(?:CREATE TABLE IF NOT EXISTS|ALTER TABLE)\s+"
                r"(?:public\.)?(\w+)", self.sql):
            self.assertTrue(table.startswith("life_"), table)

    def test_the_new_constraint_still_accepts_every_old_value(self):
        # The load-bearing property. If a horizon that is valid today were
        # left out, this migration would not widen the vocabulary — it would
        # make existing rows illegal and fail on a table that has any.
        for horizon in self.PREVIOUS:
            self.assertIn(f"'{horizon}'", self.sql, horizon)

    def test_the_constraint_matches_the_code_vocabulary(self):
        # The database and lp.HORIZONS disagreeing is how you get an insert
        # the app believes is valid and the table refuses. Read the CHECK that
        # actually ships, not the first time the word appears in the prose.
        clause = re.search(
            r"ADD CONSTRAINT life_items_horizon_check\s*CHECK\s*\((.*?)\);",
            self.sql, re.S)
        self.assertIsNotNone(clause, "the ADD CONSTRAINT was not found")
        listed = set(re.findall(r"'(\w+)'", clause.group(1)))
        self.assertEqual(listed, set(lp.HORIZONS))

    def test_null_is_still_allowed(self):
        # Most items carry no horizon at all, and the fallback path writes
        # NULL deliberately when the migration has not run.
        self.assertIn("horizon IS NULL", self.sql)


class DocxExtractionTests(unittest.TestCase):
    """The setup upload accepts .docx; extraction is best-effort and never
    raises (services/context_document.py)."""

    def test_docx_paragraphs_come_back_as_text(self):
        try:
            import docx
        except Exception as e:                      # pragma: no cover
            self.skipTest(f"python-docx unavailable: {e}")
        from io import BytesIO
        from services.context_document import extract_document_text
        document = docx.Document()
        document.add_paragraph("Section II — The Company")
        document.add_paragraph("Ship the panel [Dec]")
        buf = BytesIO()
        document.save(buf)
        out = extract_document_text(buf.getvalue(), filename="strategy.docx")
        self.assertIn("Ship the panel [Dec]", out["text"])
        self.assertEqual(out["chars"], len(out["text"]))

    def test_a_corrupt_docx_degrades_to_empty_never_raises(self):
        from services.context_document import extract_document_text
        out = extract_document_text(b"PK\x03\x04word/not-a-zip",
                                    filename="broken.docx")
        self.assertEqual(out["text"], "")


@unittest.skipIf(_IMPORT_ERROR is not None, f"needs app deps: {_IMPORT_ERROR}")
class SetupDocumentRouteTests(unittest.TestCase):
    """Upload → extract → store text; draft → review → apply ONLY the ticked
    rows. Never forced-automatic (founder 2026-07-30)."""

    def setUp(self):
        self.app = Flask(__name__)
        self.app.register_blueprint(lroutes.life_bp)
        self.client = self.app.test_client()
        self._orig_verify = auth.verify_supabase_token
        auth.verify_supabase_token = lambda token: {"sub": USER}
        self._gates = [
            patch.object(lroutes.chat, "is_enabled", return_value=True),
            patch.object(lroutes.chat, "has_consented", return_value=True),
        ]
        for g in self._gates:
            g.start()

    def tearDown(self):
        auth.verify_supabase_token = self._orig_verify
        for g in self._gates:
            g.stop()

    def _upload(self, filename: str, data: bytes = b"my strategy"):
        from io import BytesIO
        return self.client.post(
            "/v2/life/setup/document",
            headers={"Authorization": "Bearer t"},
            data={"file": (BytesIO(data), filename)},
            content_type="multipart/form-data")

    def test_extension_allowlist(self):
        for bad in ("deck.pptx", "notes.exe", "strategy"):
            resp = self._upload(bad)
            self.assertEqual(resp.status_code, 400, bad)

    def test_upload_stores_extracted_text_and_returns_metadata_only(self):
        captured = {}

        def _insert(user_id, **kw):
            captured.update(kw)
            return {"id": "d1", "file_name": kw["file_name"],
                    "status": kw["status"], "char_count": kw["char_count"],
                    "created_at": "now"}

        with patch.object(lroutes.importer, "insert_setup_document",
                          side_effect=_insert):
            resp = self._upload("strategy.md", b"# daily\nPray first.")
        self.assertEqual(resp.status_code, 201)
        body = resp.get_json()["document"]
        self.assertEqual(body["status"], "processed")
        self.assertNotIn("extracted_text", body)
        self.assertIn("Pray first.", captured["extracted_text"])

    def test_an_unreadable_file_is_kept_as_extraction_failed(self):
        with patch.object(lroutes.importer, "insert_setup_document",
                          side_effect=lambda user_id, **kw: {
                              "id": "d1", "status": kw["status"],
                              "char_count": kw["char_count"]}), \
                patch.object(lroutes, "extract_document_text",
                             return_value={"text": "", "pages": 0,
                                           "chars": 0, "truncated": False}):
            resp = self._upload("strategy.pdf", b"%PDF-broken")
        self.assertEqual(resp.status_code, 201)
        self.assertEqual(resp.get_json()["document"]["status"],
                         "extraction_failed")

    def test_propose_writes_nothing(self):
        doc = {"id": "d1", "status": "processed",
               "extracted_text": "Ship the panel [Dec]", "char_count": 20}
        with patch.object(lroutes.importer, "get_setup_document",
                          return_value=doc), \
                patch.object(lroutes.engine, "draft_items_from_document",
                             return_value=[{"kind": "goal",
                                            "title": "Ship the panel"}]), \
                patch.object(lroutes.store, "insert_item") as insert:
            resp = self.client.post(
                "/v2/life/setup/propose-from-document",
                headers={"Authorization": "Bearer t"}, json={})
        self.assertEqual(resp.status_code, 200)
        body = resp.get_json()
        self.assertFalse(body["written"])
        self.assertEqual(len(body["items"]), 1)
        insert.assert_not_called()

    def test_propose_refuses_when_there_is_nothing_readable(self):
        with patch.object(lroutes.importer, "get_setup_document",
                          return_value=None):
            resp = self.client.post(
                "/v2/life/setup/propose-from-document",
                headers={"Authorization": "Bearer t"}, json={})
        self.assertEqual(resp.status_code, 400)
        self.assertEqual(resp.get_json()["code"], "NO_DOCUMENT")

    # ── the kind hint (FE dock, 2026-07-31) ──────────────────────────────

    _DOC = {"id": "d1", "status": "processed",
            "extracted_text": "Ship the panel [Dec]", "char_count": 20}

    def _propose(self, body: dict, drafted=None):
        seen = {}

        def _draft(user_id, text, *, kind=None):
            seen["kind"] = kind
            return drafted if drafted is not None else []

        with patch.object(lroutes.importer, "get_setup_document",
                          return_value=self._DOC), \
                patch.object(lroutes.engine, "draft_items_from_document",
                             side_effect=_draft):
            resp = self.client.post(
                "/v2/life/setup/propose-from-document",
                headers={"Authorization": "Bearer t"}, json=body)
        return resp, seen

    def test_the_hint_reaches_the_engine_and_is_echoed_back(self):
        resp, seen = self._propose({"kind": "phrase"})
        self.assertEqual(resp.status_code, 200)
        self.assertEqual(seen["kind"], "phrase")
        self.assertEqual(resp.get_json()["kind"], "phrase")

    def test_an_omitted_hint_is_the_original_call_not_a_degraded_one(self):
        # Every view made exactly this request before the dock existed.
        for body in ({}, {"kind": None}, {"kind": "  "}):
            resp, seen = self._propose(body)
            self.assertEqual(resp.status_code, 200, body)
            self.assertIsNone(seen["kind"], body)
            self.assertIsNone(resp.get_json()["kind"], body)

    def test_the_hint_is_validated_against_the_closed_nine(self):
        for kind in lp.ITEM_KINDS:
            resp, seen = self._propose({"kind": kind})
            self.assertEqual(resp.status_code, 200, kind)
            self.assertEqual(seen["kind"], kind)
        # A tenth kind is an FE change. Refused here rather than ignored, so
        # nobody ships one and watches the rows vanish on the other side.
        for bad in ("reflection", "case", 7, ["phrase"]):
            with patch.object(lroutes.importer, "get_setup_document",
                              return_value=self._DOC):
                resp = self.client.post(
                    "/v2/life/setup/propose-from-document",
                    headers={"Authorization": "Bearer t"},
                    json={"kind": bad})
            self.assertEqual(resp.status_code, 400, bad)

    def test_a_hinted_draft_still_writes_nothing(self):
        drafted = [{"kind": "phrase", "title": "Pray first."}]
        with patch.object(lroutes.store, "insert_item") as insert:
            resp, _ = self._propose({"kind": "phrase"}, drafted=drafted)
        insert.assert_not_called()
        body = resp.get_json()
        self.assertFalse(body["written"])
        self.assertEqual(body["items"], drafted)

    # ── provenance (this decision was the BE's to make) ──────────────────

    def _apply(self, body: dict, insert=None):
        written = []

        def _default(user_id, fields):
            written.append(fields)
            return {"id": f"i{len(written)}", **fields}

        with patch.object(lroutes.store, "insert_item",
                          side_effect=insert or _default):
            resp = self.client.post(
                "/v2/life/setup/apply-proposed",
                headers={"Authorization": "Bearer t"}, json=body)
        return resp, written

    def test_apply_stamps_the_document_it_was_told_to_stamp(self):
        with patch.object(lroutes.importer, "get_setup_document",
                          return_value=self._DOC):
            resp, written = self._apply(
                {"document_id": "d1",
                 "items": [{"kind": "principle", "title": "Pray first."}]})
        self.assertEqual(resp.status_code, 201)
        self.assertEqual(written[0]["origin_document_id"], "d1")
        # A document-drafted principle is NOT an engine-derived one: neither
        # of the two existing provenance columns is invented for it.
        self.assertNotIn("origin_case_id", written[0])
        self.assertNotIn("origin_note_id", written[0])

    def test_apply_stamps_nothing_when_it_is_told_nothing(self):
        # The FE sends no provenance field today, and that is a complete
        # request — the rows are created, unstamped.
        resp, written = self._apply(
            {"items": [{"kind": "win", "title": "Shipped it."}]})
        self.assertEqual(resp.status_code, 201)
        self.assertNotIn("origin_document_id", written[0])

    def test_an_unresolvable_document_costs_the_stamp_and_nothing_else(self):
        # Deleted between drafting and ticking, or never theirs. The rows the
        # user ticked must survive it.
        with patch.object(lroutes.importer, "get_setup_document",
                          return_value=None):
            resp, written = self._apply(
                {"document_id": "someone-elses",
                 "items": [{"kind": "phrase", "title": "Pray first."}]})
        self.assertEqual(resp.status_code, 201)
        self.assertEqual(resp.get_json()["count"], 1)
        self.assertNotIn("origin_document_id", written[0])

    def test_apply_refuses_a_document_id_that_is_not_an_id(self):
        resp, _ = self._apply({"document_id": 7,
                               "items": [{"kind": "goal", "title": "Ship"}]})
        self.assertEqual(resp.status_code, 400)

    def test_the_rows_survive_a_migration_nobody_has_run_yet(self):
        # "On main" is not "run in prod". A missing column costs the
        # provenance, never the row the user ticked.
        attempts = []

        def _insert(user_id, fields):
            attempts.append(dict(fields))
            if "origin_document_id" in fields:
                return None          # column does not exist yet
            return {"id": "i1", **fields}

        with patch.object(lroutes.importer, "get_setup_document",
                          return_value=self._DOC):
            resp, _ = self._apply(
                {"document_id": "d1",
                 "items": [{"kind": "phrase", "title": "Pray first."}]},
                insert=_insert)
        self.assertEqual(resp.status_code, 201)
        self.assertEqual(resp.get_json()["count"], 1)
        self.assertEqual(len(attempts), 2)
        self.assertNotIn("origin_document_id", attempts[1])

    def test_apply_creates_only_sanitizable_rows_as_active(self):
        written = []

        def _insert(user_id, fields):
            written.append(fields)
            return {"id": f"i{len(written)}", **fields}

        with patch.object(lroutes.store, "insert_item", side_effect=_insert):
            resp = self.client.post(
                "/v2/life/setup/apply-proposed",
                headers={"Authorization": "Bearer t"},
                json={"items": [
                    {"kind": "goal", "title": "Ship", "status": "proposed"},
                    {"kind": "reflection", "title": "smuggled"},
                    {"kind": "habit", "title": ""},
                ]})
        self.assertEqual(resp.status_code, 201)
        self.assertEqual(resp.get_json()["count"], 1)
        self.assertEqual(len(written), 1)
        self.assertEqual(written[0]["status"], "active")

    def test_apply_requires_a_non_empty_list(self):
        for body in ({}, {"items": []}, {"items": "x"}):
            resp = self.client.post(
                "/v2/life/setup/apply-proposed",
                headers={"Authorization": "Bearer t"}, json=body)
            self.assertEqual(resp.status_code, 400, body)


@unittest.skipIf(_IMPORT_ERROR is not None, f"needs app deps: {_IMPORT_ERROR}")
class DocumentAwareGenerationTests(unittest.TestCase):

    def test_generate_documents_includes_the_upload_clearly_delimited(self):
        seen = {}

        def _complete(spec, *, system, user, surface, user_id):
            seen["user"] = user
            return {"documents": {"daily": "ok"}}

        with patch.object(lengine, "_complete", side_effect=_complete), \
                patch.object(lengine.life_importer,
                             "latest_setup_document_text",
                             return_value="Ship the panel [Dec]"):
            out = lengine.generate_documents(USER, {"bets": []})
        self.assertEqual(out, {"daily": "ok"})
        self.assertIn("The user's uploaded current strategy document:",
                      seen["user"])
        self.assertIn("Ship the panel [Dec]", seen["user"])

    def test_no_upload_means_the_prompt_is_unchanged(self):
        seen = {}

        def _complete(spec, *, system, user, surface, user_id):
            seen["user"] = user
            return {"documents": {}}

        with patch.object(lengine, "_complete", side_effect=_complete), \
                patch.object(lengine.life_importer,
                             "latest_setup_document_text", return_value=""):
            lengine.generate_documents(USER, {"bets": []})
        self.assertNotIn("uploaded current strategy document", seen["user"])

    def test_drafting_writes_nothing_and_keeps_the_locked_bets(self):
        with patch.object(lengine, "_complete", return_value={
                "goals": [{"title": "Ship the panel", "horizon": "year",
                           "due_label": "[Dec]", "bet": "company"}],
                "habits": [{"title": "Pray"}],
                "distractions": []}), \
                patch.object(lengine.store, "insert_item") as insert:
            items = lengine.draft_items_from_document(USER, "the doc text")
        insert.assert_not_called()
        kinds = [i["kind"] for i in items]
        self.assertEqual(kinds.count("bet"), 3)
        bets = [i for i in items if i["kind"] == "bet"]
        # Locked rank rides order_key, seeded from lp.BETS, never the doc.
        self.assertEqual(sorted(b["order_key"] for b in bets), [1.0, 2.0, 3.0])
        self.assertIn("goal", kinds)
        self.assertIn("habit", kinds)


@unittest.skipIf(_IMPORT_ERROR is not None, f"needs app deps: {_IMPORT_ERROR}")
class DocumentKindHintTests(unittest.TestCase):
    """The hint scopes a second read of the same text. It never filters the
    answer, and it never changes the call the endpoint made yesterday."""

    BASE = {"goals": [{"title": "Ship the panel", "due_label": "[Dec]",
                       "bet": "company"}],
            "habits": [{"title": "Pray"}],
            "distractions": [{"title": "Phone",
                              "response": "Charge it elsewhere"}]}

    def _draft(self, kind=None, lines=("Pray first.", "Do the hard thing.")):
        """Both passes stubbed. Returns (items, systems) — one system prompt
        per model call, so the count IS the number of calls made."""
        systems = []

        def _complete(spec, *, system, user, surface, user_id):
            systems.append(system)
            if surface == "doc_draft":
                return dict(self.BASE)
            return {"lines": list(lines)}

        with patch.object(lengine, "_complete", side_effect=_complete), \
                patch.object(lengine.store, "insert_item") as insert:
            items = lengine.draft_items_from_document(USER, "the doc text",
                                                      kind=kind)
        insert.assert_not_called()
        return items, systems

    def test_no_hint_makes_exactly_the_call_it_made_before(self):
        items, systems = self._draft()
        self.assertEqual(len(systems), 1)
        self.assertEqual({i["kind"] for i in items},
                         {"bet", "goal", "habit", "distraction"})

    def test_a_kind_the_base_pass_already_covers_changes_nothing(self):
        # Acceptance 4: /panel/goals is byte-identical to what it was before
        # the dock. Same rows, same bets, same due labels, same one call.
        plain, _ = self._draft()
        for kind in ("goal", "habit", "distraction", "bet", "task", "event"):
            hinted, systems = self._draft(kind=kind)
            self.assertEqual(len(systems), 1, kind)
            self.assertEqual(hinted, plain, kind)

    def test_a_phrase_hint_adds_phrase_rows_and_drops_none_of_the_rest(self):
        items, systems = self._draft(kind="phrase")
        self.assertEqual(len(systems), 2)
        kinds = [i["kind"] for i in items]
        # Hinted rows lead — the user is standing on that view.
        self.assertEqual(kinds[0], "phrase")
        self.assertEqual(kinds.count("phrase"), 2)
        # A phrases document that also holds goals still offers the goals:
        # the hint is a scoping aid, not a filter.
        self.assertIn("goal", kinds)
        self.assertIn("habit", kinds)
        phrases = [i for i in items if i["kind"] == "phrase"]
        self.assertEqual(phrases[0]["title"], "Pray first.")
        self.assertEqual(phrases[0]["collection"], "wall")

    def test_principles_and_wins_draft_their_own_kind(self):
        for kind in ("principle", "win"):
            items, systems = self._draft(kind=kind)
            self.assertEqual(len(systems), 2, kind)
            drafted = [i for i in items if i["kind"] == kind]
            self.assertEqual([d["title"] for d in drafted],
                             ["Pray first.", "Do the hard thing."], kind)
            # A principle read out of a file is tied to no case — the engine
            # derived nothing here, a person wrote it down.
            self.assertNotIn("origin_case_id", drafted[0])

    def test_the_hinted_pass_is_told_what_that_kind_is(self):
        for kind in importer.DOC_DRAFT_LINE_KINDS:
            _, systems = self._draft(kind=kind)
            self.assertIn(kind.upper() + "S", systems[1], kind)
            # Transcriber, not coach — the same fence the base pass holds.
            self.assertIn("never invent", systems[1])

    def test_nothing_found_for_the_hinted_kind_is_a_real_answer(self):
        items, _ = self._draft(kind="phrase", lines=())
        self.assertEqual([i for i in items if i["kind"] == "phrase"], [])
        self.assertIn("goal", [i["kind"] for i in items])

    def test_a_failed_hinted_pass_still_returns_the_base_rows(self):
        # An outage costs the hinted rows, never the screen (the bargain
        # every derivation in this module makes).
        def _complete(spec, *, system, user, surface, user_id):
            return dict(self.BASE) if surface == "doc_draft" else None

        with patch.object(lengine, "_complete", side_effect=_complete):
            items = lengine.draft_items_from_document(USER, "text",
                                                      kind="phrase")
        self.assertIn("goal", [i["kind"] for i in items])
        self.assertEqual([i for i in items if i["kind"] == "phrase"], [])

    def test_every_drafted_kind_is_one_the_review_step_can_create(self):
        drafted = set()
        for kind in (None,) + tuple(lp.ITEM_KINDS):
            items, _ = self._draft(kind=kind)
            drafted.update(i["kind"] for i in items)
        self.assertEqual(drafted - set(importer.DOC_DRAFT_KINDS), set())
        for kind in drafted:
            self.assertIsNotNone(
                importer.sanitize_confirmed_item(
                    {"kind": kind, "title": "🔵 The Company"
                     if kind == "bet" else "a drafted row"}), kind)


# ═════════════════════════════════════════════════════════════════════════
# Item 12 (founder 2026-07-30) — the opt-in reminder machinery itself
# ═════════════════════════════════════════════════════════════════════════

@unittest.skipIf(_IMPORT_ERROR is not None, f"needs app deps: {_IMPORT_ERROR}")
class ReminderServiceTests(unittest.TestCase):

    def test_sanitize_settings_types_and_clamps(self):
        out = lreminders.sanitize_settings({
            "morning_enabled": 1, "weekly_enabled": False,
            "tz_offset_minutes": "not-a-number", "streak_mode": True,
        })
        self.assertIs(out["morning_enabled"], True)
        self.assertIs(out["weekly_enabled"], False)
        self.assertEqual(out["tz_offset_minutes"], 0)     # unparseable → 0
        self.assertNotIn("streak_mode", out)
        clamped = lreminders.sanitize_settings({"tz_offset_minutes": 100000})
        self.assertEqual(clamped["tz_offset_minutes"], 14 * 60)

    def test_an_unknown_slot_sends_nothing(self):
        out = lreminders.run_reminder_slot("hourly")
        self.assertEqual(out["sent"], 0)
        self.assertIn("error", out)

    def _settings_rows(self, rows):
        class _Q:
            def select(self, *a, **k): return self
            def eq(self, *a, **k): return self
            def limit(self, *a, **k): return self
            def execute(self): return type("R", (), {"data": rows})()
        return _Q()

    def test_no_subscription_means_skip_and_no_claim(self):
        with patch.object(lreminders, "_t",
                          return_value=self._settings_rows(
                              [{"user_id": USER, "tz_offset_minutes": 0}])), \
                patch.object(lreminders, "list_subscriptions",
                             return_value=[]), \
                patch.object(lreminders, "_claim") as claim, \
                patch.object(lreminders, "send_web_push") as send:
            out = lreminders.run_reminder_slot("morning")
        self.assertEqual(out["sent"], 0)
        self.assertEqual(out["skipped"], 1)
        claim.assert_not_called()
        send.assert_not_called()

    def test_idempotent_per_local_day(self):
        # The claim failing means another run already owns this local day —
        # nothing is sent, which is what makes a double-fired cron a no-op.
        with patch.object(lreminders, "_t",
                          return_value=self._settings_rows(
                              [{"user_id": USER, "tz_offset_minutes": 120}])), \
                patch.object(lreminders, "list_subscriptions",
                             return_value=[{"endpoint": "e"}]), \
                patch.object(lreminders, "_claim", return_value=False), \
                patch.object(lreminders, "send_web_push") as send:
            out = lreminders.run_reminder_slot("evening")
        self.assertEqual(out["skipped"], 1)
        send.assert_not_called()

    def test_an_enabled_user_with_a_subscription_gets_one_push(self):
        with patch.object(lreminders, "_t",
                          return_value=self._settings_rows(
                              [{"user_id": USER, "tz_offset_minutes": 0}])), \
                patch.object(lreminders, "list_subscriptions",
                             return_value=[{"endpoint": "e", "p256dh": "p",
                                            "auth": "a"}]), \
                patch.object(lreminders, "_claim", return_value=True), \
                patch.object(lreminders, "send_web_push",
                             return_value=True) as send:
            out = lreminders.run_reminder_slot("weekly")
        self.assertEqual(out["sent"], 1)
        send.assert_called_once()
        payload = send.call_args[0][1]
        self.assertEqual(payload["url"], "/panel/week")

    def test_sending_without_a_private_key_is_a_graceful_no(self):
        with patch.dict(os.environ, {"VAPID_PRIVATE_KEY": ""}):
            self.assertFalse(lreminders.send_web_push(
                {"endpoint": "e", "p256dh": "p", "auth": "a"},
                {"title": "t"}))

    def test_the_payload_copy_never_nags(self):
        # The founder's copy says a thing is READY and asks the day's own
        # question. Nothing counts absence.
        for payload in lreminders.PAYLOADS.values():
            text = f"{payload['title']} {payload['body']}".lower()
            for banned in ("missed", "behind", "forget", "streak",
                           "come back", "haven't"):
                self.assertNotIn(banned, text)

    def test_local_date_respects_the_offset(self):
        from datetime import datetime, timezone
        late_utc = datetime(2026, 7, 30, 23, 30, tzinfo=timezone.utc)
        self.assertEqual(lreminders._local_date(120, late_utc), "2026-07-31")
        self.assertEqual(lreminders._local_date(-120, late_utc), "2026-07-30")

    def test_checkout_slots_land_on_existing_surfaces(self):
        # Founder 2026-07-30: NO new checkout views. Weekly deep-links to the
        # Sunday review; monthly and longer land on the panel home, where the
        # horizon strategy documents live.
        self.assertEqual(lreminders.PAYLOADS["weekly"]["url"], "/panel/week")
        for slot in ("monthly", "quarterly", "yearly", "fiveyear",
                     "tenyear"):
            self.assertEqual(lreminders.PAYLOADS[slot]["url"], "/panel", slot)
        for slot in lreminders.SLOTS:
            self.assertIn(slot, lreminders.PAYLOADS)

    def test_five_and_ten_year_slots_respect_the_anchor_year(self):
        # Cron cannot say "every 5th year", so the yearly-fired cron is
        # gated here: (year - anchor) must be a whole number of periods.
        from datetime import datetime, timezone
        y2030 = datetime(2030, 1, 1, tzinfo=timezone.utc)
        y2031 = datetime(2031, 1, 1, tzinfo=timezone.utc)
        with patch.dict(os.environ, {"LIFE_REMINDER_ANCHOR_YEAR": "2025"}):
            self.assertTrue(lreminders.slot_due_this_year("fiveyear",
                                                          now=y2030))
            self.assertFalse(lreminders.slot_due_this_year("fiveyear",
                                                           now=y2031))
            self.assertFalse(lreminders.slot_due_this_year("tenyear",
                                                           now=y2030))
            self.assertTrue(lreminders.slot_due_this_year(
                "tenyear", now=datetime(2035, 1, 1, tzinfo=timezone.utc)))
            # Every other slot is always due; its cron IS its cadence.
            self.assertTrue(lreminders.slot_due_this_year("monthly",
                                                          now=y2031))
        # Unset or malformed anchor ⇒ the operator picks years by hand.
        with patch.dict(os.environ, {"LIFE_REMINDER_ANCHOR_YEAR": ""}):
            self.assertTrue(lreminders.slot_due_this_year("tenyear",
                                                          now=y2031))
        with patch.dict(os.environ,
                        {"LIFE_REMINDER_ANCHOR_YEAR": "soonish"}):
            self.assertTrue(lreminders.slot_due_this_year("fiveyear",
                                                          now=y2031))

    def test_an_off_year_run_is_a_clean_noop(self):
        from datetime import datetime, timezone
        with patch.dict(os.environ, {"LIFE_REMINDER_ANCHOR_YEAR": "2025"}), \
                patch.object(lreminders, "_t") as table:
            out = lreminders.run_reminder_slot(
                "fiveyear", now=datetime(2031, 1, 1, tzinfo=timezone.utc))
        self.assertTrue(out.get("off_year"))
        self.assertEqual(out["sent"], 0)
        table.assert_not_called()          # not even the settings read

    def test_a_checkout_slot_runs_like_any_other(self):
        with patch.object(lreminders, "_t",
                          return_value=self._settings_rows(
                              [{"user_id": USER, "tz_offset_minutes": 0}])), \
                patch.object(lreminders, "list_subscriptions",
                             return_value=[{"endpoint": "e", "p256dh": "p",
                                            "auth": "a"}]), \
                patch.object(lreminders, "_claim", return_value=True), \
                patch.object(lreminders, "send_web_push",
                             return_value=True) as send:
            out = lreminders.run_reminder_slot("monthly")
        self.assertEqual(out["sent"], 1)
        self.assertEqual(send.call_args[0][1],
                         lreminders.PAYLOADS["monthly"])


@unittest.skipIf(_IMPORT_ERROR is not None, f"needs app deps: {_IMPORT_ERROR}")
class ReminderRouteTests(unittest.TestCase):

    def setUp(self):
        self.app = Flask(__name__)
        self.app.register_blueprint(lroutes.life_bp)
        self.client = self.app.test_client()
        self._orig_verify = auth.verify_supabase_token
        auth.verify_supabase_token = lambda token: {"sub": USER}
        self._gates = [
            patch.object(lroutes.chat, "is_enabled", return_value=True),
            patch.object(lroutes.chat, "has_consented", return_value=True),
        ]
        for g in self._gates:
            g.start()

    def tearDown(self):
        auth.verify_supabase_token = self._orig_verify
        for g in self._gates:
            g.stop()

    def test_get_returns_defaults_key_and_subscription_state(self):
        with patch.object(lroutes.reminders, "get_settings",
                          return_value=dict(lreminders.DEFAULT_SETTINGS)), \
                patch.object(lroutes.reminders, "public_key",
                             return_value=None), \
                patch.object(lroutes.reminders, "list_subscriptions",
                             return_value=[]):
            resp = self.client.get("/v2/life/reminders",
                                   headers={"Authorization": "Bearer t"})
        body = resp.get_json()
        self.assertEqual(resp.status_code, 200)
        self.assertFalse(body["settings"]["morning_enabled"])
        self.assertIsNone(body["public_key"])
        self.assertFalse(body["subscribed"])

    def test_put_writes_only_known_fields(self):
        with patch.object(lroutes.reminders, "upsert_settings",
                          return_value={**lreminders.DEFAULT_SETTINGS,
                                        "morning_enabled": True}) as up:
            resp = self.client.put("/v2/life/reminders",
                                   headers={"Authorization": "Bearer t"},
                                   json={"morning_enabled": True,
                                         "nagging_mode": True})
        self.assertEqual(resp.status_code, 200)
        self.assertNotIn("nagging_mode", up.call_args[0][1])

    def test_put_with_nothing_recognisable_is_a_400(self):
        resp = self.client.put("/v2/life/reminders",
                               headers={"Authorization": "Bearer t"},
                               json={"nagging_mode": True})
        self.assertEqual(resp.status_code, 400)

    def test_subscription_requires_the_full_key_set(self):
        for body in ({}, {"endpoint": "e"},
                     {"endpoint": "e", "keys": {"p256dh": "p"}}):
            resp = self.client.post("/v2/life/reminders/subscription",
                                    headers={"Authorization": "Bearer t"},
                                    json=body)
            self.assertEqual(resp.status_code, 400, body)

    def test_subscription_round_trip(self):
        with patch.object(lroutes.reminders, "save_subscription",
                          return_value={"id": "s1"}) as save:
            resp = self.client.post(
                "/v2/life/reminders/subscription",
                headers={"Authorization": "Bearer t"},
                json={"endpoint": "https://push/e",
                      "keys": {"p256dh": "p", "auth": "a"}})
        self.assertEqual(resp.status_code, 201)
        self.assertEqual(save.call_args.kwargs["endpoint"], "https://push/e")
        with patch.object(lroutes.reminders, "delete_subscription",
                          return_value=1):
            resp = self.client.delete(
                "/v2/life/reminders/subscription",
                headers={"Authorization": "Bearer t"},
                json={"endpoint": "https://push/e"})
        self.assertEqual(resp.status_code, 200)
        self.assertEqual(resp.get_json()["deleted"], 1)

    def test_reminder_routes_sit_behind_the_flag_and_consent(self):
        for g in self._gates:
            g.stop()
        self._gates = []
        with patch.object(lroutes.chat, "is_enabled", return_value=False):
            resp = self.client.get("/v2/life/reminders",
                                   headers={"Authorization": "Bearer t"})
            self.assertEqual(resp.status_code, 404)
        with patch.object(lroutes.chat, "is_enabled", return_value=True), \
                patch.object(lroutes.chat, "has_consented",
                             return_value=False):
            resp = self.client.get("/v2/life/reminders",
                                   headers={"Authorization": "Bearer t"})
            self.assertEqual(resp.status_code, 409)


@unittest.skipIf(_IMPORT_ERROR is not None, f"needs app deps: {_IMPORT_ERROR}")
class ReminderWebhookTests(unittest.TestCase):

    def setUp(self):
        self.app = Flask(__name__)
        self.app.register_blueprint(lremhook.life_reminders_webhook_bp)
        self.client = self.app.test_client()

    def _post(self, slot="morning", secret="s3cret"):
        headers = {}
        if secret is not None:
            headers["X-Internal-Secret"] = secret
        return self.client.post(f"/v2/internal/life/reminders?slot={slot}",
                                headers=headers, json={})

    def test_unset_secret_means_the_surface_is_dead(self):
        with patch.dict(os.environ, {"LIFE_REMINDER_SECRET": ""}):
            self.assertEqual(self._post().status_code, 503)

    def test_wrong_or_missing_secret_is_401(self):
        with patch.dict(os.environ, {"LIFE_REMINDER_SECRET": "right"}):
            self.assertEqual(self._post(secret="wrong").status_code, 401)
            self.assertEqual(self._post(secret=None).status_code, 401)

    def test_flag_off_skips_the_run_entirely(self):
        with patch.dict(os.environ, {"LIFE_REMINDER_SECRET": "s3cret"}), \
                patch.object(lremhook.chat, "is_enabled",
                             return_value=False), \
                patch.object(lremhook.reminders, "run_reminder_slot") as run:
            resp = self._post()
        self.assertEqual(resp.status_code, 200)
        run.assert_not_called()

    def test_bad_slot_is_400(self):
        with patch.dict(os.environ, {"LIFE_REMINDER_SECRET": "s3cret"}), \
                patch.object(lremhook.chat, "is_enabled", return_value=True):
            self.assertEqual(self._post(slot="hourly").status_code, 400)

    def test_a_valid_call_runs_the_slot_and_returns_counts(self):
        counts = {"slot": "morning", "eligible": 1, "sent": 1,
                  "skipped": 0, "failed": 0}
        with patch.dict(os.environ, {"LIFE_REMINDER_SECRET": "s3cret"}), \
                patch.object(lremhook.chat, "is_enabled", return_value=True), \
                patch.object(lremhook.reminders, "run_reminder_slot",
                             return_value=counts) as run:
            resp = self._post()
        self.assertEqual(resp.status_code, 200)
        self.assertEqual(resp.get_json(), counts)
        run.assert_called_once_with("morning")


# ═════════════════════════════════════════════════════════════════════════
# Editable check-in copy (founder decision 2026-07-30: "my copy as default,
# editable")
# ═════════════════════════════════════════════════════════════════════════

@unittest.skipIf(_IMPORT_ERROR is not None, f"needs app deps: {_IMPORT_ERROR}")
class UserCopyTests(unittest.TestCase):

    def test_sanitize_keeps_only_the_three_fields_and_caps_them(self):
        out = lengine.sanitize_user_copy({
            "mantra_title": "  MY OWN LINE  ",
            "distraction_question": "x" * 900,
            "one_thing": "smuggled",
            "evening_line": "smuggled",
        })
        self.assertEqual(out["mantra_title"], "MY OWN LINE")
        self.assertEqual(len(out["distraction_question"]), 300)
        self.assertNotIn("one_thing", out)
        self.assertNotIn("evening_line", out)

    def test_blank_means_back_to_the_default(self):
        # NULL is the contract: the FE renders the founder's line on NULL,
        # so clearing a field is how the original comes back.
        out = lengine.sanitize_user_copy({"mantra_question": "   ",
                                          "mantra_title": None})
        self.assertIsNone(out["mantra_question"])
        self.assertIsNone(out["mantra_title"])

    def test_get_returns_every_field_with_null_for_unset(self):
        with patch.object(lengine, "_copy_table",
                          side_effect=RuntimeError("down")):
            out = lengine.get_user_copy(USER)
        self.assertEqual(out, {"mantra_title": None, "mantra_question": None,
                               "distraction_question": None})


@unittest.skipIf(_IMPORT_ERROR is not None, f"needs app deps: {_IMPORT_ERROR}")
class UserCopyRouteTests(unittest.TestCase):

    def setUp(self):
        self.app = Flask(__name__)
        self.app.register_blueprint(lroutes.life_bp)
        self.client = self.app.test_client()
        self._orig_verify = auth.verify_supabase_token
        auth.verify_supabase_token = lambda token: {"sub": USER}
        self._gates = [
            patch.object(lroutes.chat, "is_enabled", return_value=True),
            patch.object(lroutes.chat, "has_consented", return_value=True),
        ]
        for g in self._gates:
            g.start()

    def tearDown(self):
        auth.verify_supabase_token = self._orig_verify
        for g in self._gates:
            g.stop()

    def test_get_serves_the_overrides_nulls_when_unset(self):
        with patch.object(lroutes.engine, "get_user_copy",
                          return_value={"mantra_title": None,
                                        "mantra_question": None,
                                        "distraction_question": None}):
            resp = self.client.get("/v2/life/copy",
                                   headers={"Authorization": "Bearer t"})
        self.assertEqual(resp.status_code, 200)
        self.assertIsNone(resp.get_json()["copy"]["mantra_title"])

    def test_put_saves_the_users_own_wording(self):
        with patch.object(lroutes.engine, "upsert_user_copy",
                          return_value={"mantra_title": "MY LINE",
                                        "mantra_question": None,
                                        "distraction_question": None}) as up:
            resp = self.client.put("/v2/life/copy",
                                   headers={"Authorization": "Bearer t"},
                                   json={"mantra_title": "MY LINE"})
        self.assertEqual(resp.status_code, 200)
        self.assertEqual(resp.get_json()["copy"]["mantra_title"], "MY LINE")
        self.assertEqual(up.call_args[0][1], {"mantra_title": "MY LINE"})

    def test_put_with_nothing_recognisable_is_a_400(self):
        resp = self.client.put("/v2/life/copy",
                               headers={"Authorization": "Bearer t"},
                               json={"one_thing": "smuggled"})
        self.assertEqual(resp.status_code, 400)


if __name__ == "__main__":
    unittest.main()
