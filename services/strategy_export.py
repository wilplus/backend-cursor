"""The life-panel strategy as a real downloadable document (founder
2026-07-27, backlog story #11).

`GET /v2/life/strategy/download` served JSON, so "download" meant the FE
hand-rolling a .txt. Here the same assembled document comes back as a real
PDF or a real .docx — an editable Word file was the founder's explicit ask, so
the strategy can be reviewed and shared offline.

Both renderers lazy-import their dependency and a missing/failed import
DEGRADES TO MARKDOWN rather than 500ing, exactly as
`services/dev_tasks.py::export_pdf` degrades to the ZIP. A download must never
be the thing that takes the panel down.

Content only — no scores, no verdicts, nothing derived (AC-9). This module
formats the student's own words and adds none of its own.
"""
from __future__ import annotations

import io
import logging
import re
from typing import Any

logger = logging.getLogger(__name__)

MD = ("text/markdown; charset=utf-8", ".md")
PDF = ("application/pdf", ".pdf")
DOCX = ("application/vnd.openxmlformats-officedocument"
        ".wordprocessingml.document", ".docx")

FORMATS = ("json", "md", "pdf", "docx")

_TITLE = "Your strategy"
_PARA_SPLIT = re.compile(r"\n[ \t]*\n")


def horizon_label(horizon: Any) -> str:
    """"five_year" → "Five year". The stored keys are snake_case; a document
    heading is not. Pure."""
    return str(horizon or "").replace("_", " ").strip().capitalize()


def sections(latest: dict, horizons: Any) -> list:
    """[(heading, body)] in HORIZON order — the same order and the same
    source the JSON download uses, so every format says the same thing. A
    horizon with no document, or an empty one, is skipped. Pure."""
    out = []
    for horizon in (horizons or ()):
        row = (latest or {}).get(horizon)
        body = ((row or {}).get("body") or "").strip()
        if body:
            out.append((horizon_label(horizon), body))
    return out


def to_markdown(parts: list) -> str:
    """The assembled document as markdown — also the degradation target when
    a renderer's dependency is unavailable. Pure."""
    return "\n\n".join(f"# {head}\n\n{body}" for head, body in parts)


def _render_pdf(parts: list) -> bytes:
    from xml.sax.saxutils import escape

    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.platypus import Paragraph, SimpleDocTemplate

    margin = 18 * mm
    ss = getSampleStyleSheet()
    st_title = ParagraphStyle("wl_s_title", parent=ss["Title"], fontSize=20,
                              spaceAfter=14, alignment=0)
    st_head = ParagraphStyle("wl_s_head", parent=ss["Heading2"], fontSize=13,
                             spaceBefore=16, spaceAfter=4,
                             textColor="#111214")
    st_body = ParagraphStyle("wl_s_body", parent=ss["BodyText"], fontSize=10.5,
                             leading=16, spaceAfter=6)

    flow: list[Any] = [Paragraph(_TITLE, st_title)]
    for head, body in parts:
        flow.append(Paragraph(escape(head), st_head))
        for para in _PARA_SPLIT.split(body):
            if para.strip():
                flow.append(Paragraph(
                    escape(para.strip()).replace("\n", "<br/>"), st_body))

    buf = io.BytesIO()
    SimpleDocTemplate(
        buf, pagesize=A4, leftMargin=margin, rightMargin=margin,
        topMargin=margin, bottomMargin=margin,
        title=_TITLE, author="WillpowerLab",
    ).build(flow)
    return buf.getvalue()


def _render_docx(parts: list) -> bytes:
    from docx import Document

    doc = Document()
    doc.add_heading(_TITLE, level=0)
    for head, body in parts:
        doc.add_heading(head, level=1)
        for para in _PARA_SPLIT.split(body):
            if para.strip():
                doc.add_paragraph(para.strip())
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def export(latest: dict, horizons: Any, fmt: Any) -> tuple:
    """(payload: bytes, mimetype: str, filename: str) for `fmt` in
    md/pdf/docx.

    An unknown format, an empty strategy, or a renderer whose dependency is
    missing or blows up all land on MARKDOWN — a 200 with real content beats a
    500 on a download button. The caller reads the returned mimetype rather
    than assuming one from the request (the FE contract says the same)."""
    parts = sections(latest, horizons)
    md = to_markdown(parts)
    fmt = str(fmt or "").strip().lower()
    if fmt in ("pdf", "docx") and parts:
        render, (mime, ext) = ((_render_pdf, PDF) if fmt == "pdf"
                               else (_render_docx, DOCX))
        try:
            return render(parts), mime, f"strategy{ext}"
        except Exception:  # noqa: BLE001
            logger.warning(
                "strategy_export: %s render unavailable, serving markdown",
                fmt, exc_info=True)
    return md.encode("utf-8"), MD[0], f"strategy{MD[1]}"
